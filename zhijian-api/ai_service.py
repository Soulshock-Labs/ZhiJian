"""
ai_service — AI 内容生成模块

职责：
  build_system_prompt()        — 系统角色 Prompt
  build_user_prompt()          — 用户输入 Prompt 构建（注入理念词、班级特征、模板提纲）
  generate_content()           — 调用 DashScope LLM，返回归一化内容字典
  _extract_template_outline()  — 从上传 .docx 中提取语义提纲（14 条关键词）
  _normalize_content_payload() — 强约束输出结构（防模型波动）
  _parse_json_payload()        — JSON 解析 + 容错抽取

依赖：
  word_engine.field_map — PHILOSOPHY_HINTS, CLASS_LEVEL_HINTS, ACTIVITY_LABEL_MAP
  openai SDK（OpenAI-compatible，默认指向 DashScope）
"""
from __future__ import annotations

import io
import json
import logging
import os
import re
import textwrap
from typing import Optional

from docx import Document
from fastapi import HTTPException
from openai import OpenAI

from core.settings import AI_MODEL, MOONSHOT_API_KEY, MOONSHOT_BASE_URL, OPENAI_API_KEY, OPENAI_BASE_URL, VOICE_TRANSCRIBE_MODEL
from word_engine.field_map import PHILOSOPHY_HINTS, CLASS_LEVEL_HINTS, ACTIVITY_LABEL_MAP

logger = logging.getLogger(__name__)

# ──────────────────────────────────────────────
# AI 客户端（模块加载时初始化）
# ──────────────────────────────────────────────
ALLOW_MOCK_CONTENT = str(os.getenv("ALLOW_MOCK_CONTENT", "0")).strip().lower() in {"1", "true", "yes", "on"}

client = OpenAI(api_key=MOONSHOT_API_KEY or "not-configured", base_url=MOONSHOT_BASE_URL)
voice_client = OpenAI(api_key=OPENAI_API_KEY, base_url=OPENAI_BASE_URL) if OPENAI_API_KEY else None

_ACTIVITY_KEYS = ("morning", "outdoor", "environment", "life", "study", "game", "area", "family", "departure")

_OUTLINE_KEYS = (
    "主题", "目标", "准备", "过程", "导入", "延伸", "反思", "观察",
    "评价", "家园", "晨间", "户外", "区域", "离园", "周一", "周二", "周三", "周四", "周五",
)


def _raise_if_invalid_key(exc: Exception) -> None:
    text = str(exc)
    low = text.lower()
    if "401" in text or "invalid_api_key" in low or "incorrect api key" in low:
        raise HTTPException(
            status_code=503,
            detail="AI API Key 无效或已过期，请检查 Cloud Run 环境变量 MOONSHOT_API_KEY / DEEPSEEK_API_KEY。",
        )


# ──────────────────────────────────────────────
# Prompt 构建
# ──────────────────────────────────────────────

def build_system_prompt() -> str:
    return textwrap.dedent("""
        你是一位拥有10年经验的资深幼儿园教研主任，擅长编写专业的幼儿园活动计划。
        你的输出必须：
        1. 严格按照 JSON 格式返回，不包含任何 Markdown 代码块标记
        2. 内容专业、温暖、具有可操作性
        3. 每条目标/环节内容言简意赅，控制在 80 字以内
        4. 评价部分要有具体的观察维度和记录建议
        5. 必须紧扣"教学主题"和"教育理念"，禁止输出与主题无关的泛化套话
        6. 不要编造模板中不存在的额外模块，只填充给定 JSON 结构
    """).strip()


def build_user_prompt(
    theme: str,
    phil: str,
    activities: list[str],
    child_initiative: bool,
    child_desc: str,
    template_outline: Optional[list[str]] = None,
    class_level: str = "",
) -> str:
    phil_hint = PHILOSOPHY_HINTS.get(phil, "")
    class_hint = CLASS_LEVEL_HINTS.get(class_level, "")
    acts_str = "、".join(activities) if activities else "区域活动、户外活动"
    outline_hint = ""
    if template_outline:
        outline_items = "\n".join(f"- {x}" for x in template_outline[:14])
        outline_hint = f"\n【老师模板提纲参考（优先对齐语义结构，不要求复刻版式）】\n{outline_items}\n"

    return textwrap.dedent(f"""
        请为以下幼儿园活动生成专业内容，以 JSON 格式返回。

        【输入信息】
        - 教学主题：{theme}
        - 教育理念 / 园本特色：{phil}
        - 班级：{class_level or "未指定（请按中班水平默认）"}
        - 活动重点：{acts_str}
        - 幼儿自主发起活动：{"是，" + child_desc if child_initiative and child_desc else "否"}

        【理念专业词汇要求】
        {phil_hint}

        【班级年龄特征约束】
        {class_hint or "按中班（4-5岁）水平默认生成，活动难度与目标表述取中间值。"}
        {outline_hint}
        【贴合要求（严格）】
        - 所有字段必须围绕「{theme}」，不能写成任意主题都通用的话术。
        - 优先贴合老师模板提纲语义；若提纲含"周一~周五/观察/反思"等，应在对应字段体现。
        - 周一到周五应体现五大领域（健康、语言、社会、科学、艺术）的均衡分配，不可集中在单一领域。
        - 输出应体现园本特色，不要写成空泛口号。
        - 仅返回可被 json.loads 解析的对象 JSON。

        【内容适宜性（必须遵守）】
        - 主题与活动名称必须是幼儿园教育相关的温和内容
        - 禁止军事、战争、暴力、政治、恐怖等不适宜幼儿的内容
        - 优先选择生活经验、自然、季节、社交、情感、游戏等主题
        - 活动名称要童趣、积极、安全

        【JSON 输出格式（严格遵守）】
        {{
          "weekly_targets": {{
            "teaching": "本周教学目标，1 句话，≤30字（贴合《3-6岁指南》并结合班级现状）",
            "life": "2-3 条生活目标，每条≤15字，用「能」字句，逗号或换行分隔；不要自己加 1. 2. 编号（系统会加）",
            "family": "2-3 条家园共育目标，每条≤15字，逗号或换行分隔；不要自己加 1. 2. 编号",
            "environment": "2-3 条环创目标，每条≤15字，逗号或换行分隔；不要自己加 1. 2. 编号"
          }},
          "goals": ["每条目标≤15字，简洁、动词开头", "如：能跟随节奏走跑跳", "用「能」「会」「乐意」等动词开头", "目标4（可选）", "目标5（可选）"],
          "preparation": ["材料1", "材料2", "材料3"],
          "activities": {{
            "morning":     "晨间运动建议（简洁，围绕主题，≤60字）",
            "outdoor":     "5种不同户外活动名称（每个 4-8 字）。可用逗号、换行或顿号分隔，如：螃蟹赛跑、贴人、跳房子、好玩的报纸、听哨子 或 螃蟹赛跑\\n贴人\\n跳房子\\n好玩的报纸\\n听哨子",
            "environment": "环创活动方案 2-3 条（每条≤30字）。可用逗号或换行分隔，格式自由，系统将自动转换为 1. 2. 3. 编号式",
            "life":        "生活活动目标 2-3 条（每条≤20字）。用「能」字句表达，如：能根据需求选择分量合适的点心、能独立整理好衣物。可用逗号或换行分隔",
            "study":       "集中活动（仅适用中班/大班，5 个名称，动静结合）。可用逗号或换行分隔，如：语言《故事》、体育《跳绳》，系统自动按领域分配",
            "game":        "游戏活动（5 种不同游戏，每个 4-8 字）。可用逗号或换行分隔，如：夹豆子、贴人、好玩的报纸、跳房子、听哨子",
            "area":        "区域活动建议（简洁，围绕主题，≤60字）",
            "family":      "家园共育建议 2-3 条（每条≤25字）。可用逗号或换行分隔",
            "departure":   "离园活动建议（简洁，≤40字）"
          }},
          "child_initiative_note": "如果幼儿有自主发起活动，给出教师跟进建议（无则留空，≤60字）",
          "nap_guidance": "午睡指导（睡姿、穿脱衣物、衣物摆放、睡前习惯，≤50字）",
          "evaluation": "评价与反思建议（含观察维度、记录方式、改进方向，≤80字）"
        }}

        【内容生成准则】
        ✓ 你负责生成**内容和数据**，系统负责处理**格式转换**
        ✓ outdoor/game/study：生成 5 种不同活动名称，用任意分隔符（逗号、换行、空格皆可）
        ✓ life/family/environment：生成 2-3 条建议，分隔符自由
        ✓ 系统会自动将这些转换为标准格式："周一《活动1》\\n周二《活动2》"或"1. 项目1\\n2. 项目2"
        ✓ 你只需保证：5 个不同项、字数不超限、内容围绕主题

        只返回 JSON，不要任何其他文字。
    """).strip()


# ──────────────────────────────────────────────
# 模板提纲抽取（让 LLM 内容与模板语义更一致）
# ──────────────────────────────────────────────

def _extract_template_outline(template_bytes: bytes, max_items: int = 14) -> list[str]:
    try:
        doc = Document(io.BytesIO(template_bytes))
    except Exception:
        return []

    def _clean(s: str) -> str:
        t = re.sub(r"\s+", " ", (s or "")).strip()
        return t.strip("：:;；,.。")

    rows: list[str] = []
    seen: set[str] = set()

    def push(raw: str) -> None:
        t = _clean(raw)
        if not t or t in seen or len(t) < 2 or len(t) > 26:
            return
        if not any(k in t for k in _OUTLINE_KEYS):
            return
        seen.add(t)
        rows.append(t)

    for p in doc.paragraphs:
        push(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                push(cell.text)
    return rows[:max_items]


# ──────────────────────────────────────────────
# JSON 解析 + 内容归一化
# ──────────────────────────────────────────────

def _parse_json_payload(raw_text: str) -> dict:
    raw = str(raw_text or "").strip()
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    try:
        parsed = json.loads(raw)
        return parsed if isinstance(parsed, dict) else {}
    except Exception:
        pass
    start = raw.find("{")
    end = raw.rfind("}")
    if start >= 0 and end > start:
        try:
            parsed = json.loads(raw[start:end + 1])
            return parsed if isinstance(parsed, dict) else {}
        except Exception:
            return {}
    return {}


def _normalize_content_payload(
    payload: dict,
    *,
    theme: str,
    phil: str,
    activities: list[str],
    child_initiative: bool,
    child_desc: str,
) -> dict:
    """强约束输出：即使模型波动，也保证结构齐全、字段稳定、语义一致。"""
    fallback = _mock_content(theme, phil, activities)
    data = payload if isinstance(payload, dict) else {}

    weekly_targets_in = data.get("weekly_targets", {})
    weekly_targets_in = weekly_targets_in if isinstance(weekly_targets_in, dict) else {}
    weekly_targets = {
        "teaching": str(weekly_targets_in.get("teaching", "") or "").strip(),
        "life":     str(weekly_targets_in.get("life", "") or "").strip(),
        "family":   str(weekly_targets_in.get("family", "") or "").strip(),
        "environment": str(weekly_targets_in.get("environment", "") or "").strip(),
    }
    if not weekly_targets["teaching"]:
        weekly_targets["teaching"] = f"围绕「{theme}」落实五大领域核心经验，体现{phil}特色并兼顾班级个体差异。"
    if not weekly_targets["life"]:
        weekly_targets["life"] = "聚焦卫生保健、安全规则与生活自理，形成可持续的日常习惯。"
    if not weekly_targets["family"]:
        weekly_targets["family"] = "通过家园沟通与家庭延伸任务，形成教育一致性与共同支持。"
    if not weekly_targets["environment"]:
        weekly_targets["environment"] = f"基于「{theme}」优化主题角与材料投放，支持角色扮演与探究活动。"

    goals_raw = data.get("goals", [])
    goals = [str(x).strip() for x in (goals_raw or []) if str(x).strip()]
    if len(goals) < 3:
        goals = [str(x).strip() for x in fallback["goals"]]
    goals = goals[:5]

    prep_raw = data.get("preparation", [])
    preparation = [str(x).strip() for x in (prep_raw or []) if str(x).strip()]
    if len(preparation) < 3:
        preparation = [str(x).strip() for x in fallback["preparation"]]
    preparation = preparation[:5]

    in_acts = data.get("activities", {})
    in_acts = in_acts if isinstance(in_acts, dict) else {}
    selected = set(activities or [])
    normalized_acts: dict[str, str] = {}
    for key in _ACTIVITY_KEYS:
        v = str(in_acts.get(key, "") or "").strip()
        if not v:
            label = ACTIVITY_LABEL_MAP.get(key, key)
            if key in selected:
                v = f"围绕「{theme}」开展{label}，按{phil}理念设计可执行步骤，并记录幼儿关键表现。"
            else:
                v = f"结合「{theme}」预留{label}联动建议，支持后续按班级实际灵活启用。"
        normalized_acts[key] = v

    child_note = str(data.get("child_initiative_note", "") or "").strip()
    if child_initiative:
        if not child_note:
            child_note = (
                f"围绕幼儿自主发起内容进行追问与延展：{child_desc}"
                if child_desc else
                "观察并承接幼儿自主发起线索，采用小步追问与材料支持，形成可延续的探究路径。"
            )
    else:
        child_note = ""

    evaluation = str(data.get("evaluation", "") or "").strip() or str(fallback["evaluation"])
    nap_guidance = str(data.get("nap_guidance", "") or "").strip() or str(fallback.get("nap_guidance", ""))

    return {
        "weekly_targets": weekly_targets,
        "goals": goals,
        "preparation": preparation,
        "activities": normalized_acts,
        "child_initiative_note": child_note,
        "nap_guidance": nap_guidance,
        "evaluation": evaluation,
    }


def _mock_content(theme: str, phil: str, activities: list[str]) -> dict:
    return {
        "weekly_targets": {
            "teaching": f"围绕「{theme}」在五大领域开展均衡经验建构，体现{phil}特色。",
            "life": "提升卫生保健、安全与生活自理能力，形成稳定日常习惯。",
            "family": "明确家园配合重点，促进家庭端对主题学习的延伸支持。",
            "environment": f"依据「{theme}」配置环创与材料支持，如角色区与主题操作材料。",
        },
        "goals": [
            f"【健康领域】围绕「{theme}」，发展幼儿大肌肉协调能力与身体控制能力",
            f"【语言领域】通过「{theme}」情境丰富词汇，培养表达与倾听能力",
            f"【社会领域】在「{theme}」活动中建立合作意识与规则感",
            f"【科学领域】以「{theme}」为载体，激发观察与探究兴趣",
            f"【艺术领域】感受「{theme}」之美，发展创意表达能力",
        ],
        "preparation": [
            f"「{theme}」主题相关图片卡片与实物材料",
            "《3-6岁儿童学习与发展指南》领域目标对照单",
            "幼儿观察记录表",
        ],
        "activities": {a: f"围绕「{theme}」开展{a}（{phil}理念指导）" for a in activities},
        "child_initiative_note": "",
        "nap_guidance": "睡前提醒幼儿如厕与饮水；指导仰卧或侧卧，衣物分类摆放；起床后协助整理穿脱衣物，培养自理能力。",
        "evaluation": (
            f"通过观察幼儿在「{theme}」活动中的参与度、语言表达及合作行为，"
            f"参照{phil}理念的核心经验指标进行发展性评价。"
        ),
    }


# ──────────────────────────────────────────────
# 主生成函数
# ──────────────────────────────────────────────

def generate_content(
    theme: str,
    phil: str,
    activities: list[str],
    child_initiative: bool,
    child_desc: str,
    template_outline: Optional[list[str]] = None,
    class_level: str = "",
) -> dict:
    """调用 Kimi / Moonshot，返回解析后的内容字典。"""
    if not MOONSHOT_API_KEY:
        logger.warning("MOONSHOT_API_KEY 缺失，回退 Mock 生成（/generate 旧接口）")
        return _normalize_content_payload(
            _mock_content(theme, phil, activities),
            theme=theme, phil=phil, activities=activities,
            child_initiative=child_initiative, child_desc=child_desc,
        )

    messages = [
        {"role": "system", "content": build_system_prompt()},
        {"role": "user", "content": build_user_prompt(
            theme, phil, activities, child_initiative, child_desc, template_outline,
            class_level=class_level,
        )},
    ]

    try:
        response = client.chat.completions.create(
            model=AI_MODEL, messages=messages, temperature=1, max_tokens=4096,
        )
        raw = response.choices[0].message.content.strip()
        parsed = _parse_json_payload(raw)
        if not parsed:
            retry_messages = messages + [
                {"role": "assistant", "content": raw},
                {"role": "user", "content": "你上一条输出不是可解析 JSON。请仅返回一个合法 JSON 对象，不要附加任何文字。"},
            ]
            retry_resp = client.chat.completions.create(
                model=AI_MODEL, messages=retry_messages, temperature=1, max_tokens=4096,
            )
            parsed = _parse_json_payload(retry_resp.choices[0].message.content.strip())
        if not parsed:
            raise json.JSONDecodeError("invalid json payload", raw, 0)
        return _normalize_content_payload(
            parsed, theme=theme, phil=phil, activities=activities,
            child_initiative=child_initiative, child_desc=child_desc,
        )

    except json.JSONDecodeError as e:
        if not ALLOW_MOCK_CONTENT:
            logger.error("AI JSON 解析失败且禁止 Mock 回退：%s", e)
            raise HTTPException(status_code=502, detail="AI 返回内容格式异常（JSON 解析失败），请重试或检查模型配置。")
        logger.warning("AI JSON 解析失败，ALLOW_MOCK_CONTENT=1，回退规范化 Mock：%s", e)
        return _normalize_content_payload(
            _mock_content(theme, phil, activities),
            theme=theme, phil=phil, activities=activities,
            child_initiative=child_initiative, child_desc=child_desc,
        )
    except Exception as e:
        _raise_if_invalid_key(e)
        raise HTTPException(status_code=502, detail=f"调用 AI API 失败：{e}")
