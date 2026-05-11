"""
智伴幼师 · AI幼师助手  —  FastAPI 后端
==========================================
POST /generate          → 周/月活动计划表 → 填充 Word 模板
POST /preview           → 仅返回 AI 内容 JSON（调试用）
POST /template/analyze  → 上传模板，返回与 TEMPLATE_STANDARD v1.1 对齐的自检 JSON
POST /generate-weekly   → 基于主题自动生成五天周计划 JSON
POST /generate-daily    → 将周计划某一天拆解为「导入/过程/延伸/反思」四维日教案 → 填充 Word 模板
POST /redeem            → 卡密兑换：验证、核销、开通服务

铁律：遍历 Word 表格 → 关键字匹配单元格 → 保留原样式精准填充，格式分毫不动。
依赖：fastapi  uvicorn  aspose-words  python-docx  openai  python-multipart  python-dotenv  lxml（净空引擎 kindergarten_template_cleaner）
"""

from __future__ import annotations

import asyncio
import base64
import copy
import importlib
import io
import json
import logging
import os
import re
import secrets
import string
import textwrap
import tempfile
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Optional
from urllib.parse import quote
from uuid import uuid4

from dotenv import load_dotenv
from fastapi import Body, FastAPI, File, Form, Header, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from openai import OpenAI
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml.ns import qn
from docx.shared import Pt

# 注意：不要在模块导入阶段直接 import aspose.words。
# 在部分环境中该导入会触发 CLR 级崩溃（进程退出），无法由 try/except 捕获。
aw = None  # type: ignore

from kindergarten_template_cleaner import clean_docx_bytes, find_all_day_header_rows
from redeem_routes import register_redeem_routes

# Prompt 工程系统（可复现的高质量生成）
from prompt_engineering.prompt_config import get_prompt_template

# Firestore（懒加载，Cloud Run 上有 ADC 自动认证）
try:
    from google.cloud import firestore as _firestore
    _FS_CLIENT: Optional["_firestore.Client"] = None

    def _fs() -> "_firestore.Client":
        global _FS_CLIENT
        if _FS_CLIENT is None:
            _FS_CLIENT = _firestore.Client()
        return _FS_CLIENT

    FIRESTORE_ENABLED = True
except Exception:
    FIRESTORE_ENABLED = False
    _FS_CLIENT = None

    def _fs():
        return None

load_dotenv()

logger = logging.getLogger(__name__)

# 与前端角标、导出标记一致
APP_VERSION = os.getenv("APP_VERSION", "1.2.1")


def _env_truthy(name: str, default: str = "0") -> bool:
    raw = os.getenv(name, default)
    return str(raw).strip().lower() in {"1", "true", "yes", "on"}


# 安全开关：Aspose 在部分环境会触发 CLR 级崩溃（无法由 Python try/except 捕获）。
# 默认关闭，显式配置 ENABLE_ASPOSE_WORDS=1 再启用。
ENABLE_ASPOSE_WORDS = _env_truthy("ENABLE_ASPOSE_WORDS", "0")
if not ENABLE_ASPOSE_WORDS:
    logger.warning("Aspose.Words 默认关闭：将使用 python-docx 导出（可设 ENABLE_ASPOSE_WORDS=1 开启）")
else:
    logger.info("Aspose.Words 开关已开启：将在首次导出时尝试懒加载")

# 生成内容安全开关：默认不允许静默回退 Mock，避免线上“看似成功但内容跑偏”。
ALLOW_MOCK_CONTENT = _env_truthy("ALLOW_MOCK_CONTENT", "0")


def _aw_lazy_import():
    global aw
    if aw is None:
        aw = importlib.import_module("aspose.words")  # noqa: PLC0415
    return aw

# ──────────────────────────────────────────────
# App & CORS
# ──────────────────────────────────────────────
app = FastAPI(
    title="智伴幼师 API",
    description="AI幼师助手后端服务",
    version=APP_VERSION,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # 含 file:// / 静态页跨域调 Cloud Run；勿与 allow_credentials=True 同用 *
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ──────────────────────────────────────────────
# 阿里云百炼 / OpenAI-compatible client
# ──────────────────────────────────────────────
DASHSCOPE_API_KEY = os.getenv("DASHSCOPE_API_KEY", "")
DASHSCOPE_BASE_URL = os.getenv(
    "DASHSCOPE_BASE_URL", "https://dashscope.aliyuncs.com/compatible-mode/v1"
)
AI_MODEL = os.getenv("AI_MODEL", "qwen-max")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
OPENAI_BASE_URL = os.getenv("OPENAI_BASE_URL", "https://api.openai.com/v1")
VOICE_TRANSCRIBE_MODEL = os.getenv("VOICE_TRANSCRIBE_MODEL", "whisper-1")

client = OpenAI(
    api_key=DASHSCOPE_API_KEY,
    base_url=DASHSCOPE_BASE_URL,
)
voice_client = (
    OpenAI(api_key=OPENAI_API_KEY, base_url=OPENAI_BASE_URL)
    if OPENAI_API_KEY
    else None
)


def _raise_if_invalid_dashscope_key(exc: Exception) -> None:
    """
    DashScope / OpenAI 兼容客户端在 Key 错误时通常返回 401。
    与「未配置」区分：环境变量有值但无效时，/health 仍会显示已配置，需给出可操作说明。
    """
    text = str(exc)
    low = text.lower()
    if (
        "401" in text
        or "invalid_api_key" in low
        or "incorrect api key" in low
    ):
        raise HTTPException(
            status_code=503,
            detail=(
                "阿里云百炼 API Key 无效或已过期。请在部署环境（Cloud Run 环境变量/密钥）"
                "更新 DASHSCOPE_API_KEY 为百炼控制台生成的完整密钥。"
                "使用通义 Qwen 时请保持 DASHSCOPE_BASE_URL="
                "https://dashscope.aliyuncs.com/compatible-mode/v1，"
                "勿将 DeepSeek / OpenAI 的 Key 与百炼 Base URL 混用。"
            ),
        )


# ──────────────────────────────────────────────
# 标准模板下载计数（服务端持久化）
# ──────────────────────────────────────────────
_BASE_DIR = Path(__file__).resolve().parent
_KNOWLEDGE_BASE_DIR = _BASE_DIR / "knowledge_base"
_KNOWLEDGE_INDEX_FILE = _KNOWLEDGE_BASE_DIR / "indexes" / "knowledge_index.json"
_KNOWLEDGE_ROUTE_FILE = _KNOWLEDGE_BASE_DIR / "indexes" / "profile_routes.json"
_TEMPLATE_STATS_FILE = _BASE_DIR / "template_download_stats.json"
_DEFAULT_TEMPLATE_STATS = {
    "weekly": 0,  # 周/活动计划模板
    "daily": 0,   # 日教案模板
    "cleaned": 0, # 净空版模板下载
}

_APP_STATS_FILE = _BASE_DIR / "app_public_stats.json"
_FEEDBACK_LOG_FILE = _BASE_DIR / "feedback_messages.jsonl"
_WEEKLY_DRAFT_LOG_FILE = _BASE_DIR / "weekly_draft_sessions.jsonl"
_REGISTER_LOG_FILE = _BASE_DIR / "registrations.jsonl"
_REDEEM_CODES_FILE = _BASE_DIR / "redeem_codes.json"
# 若设置 gs://bucket/path/redeem_codes.json，核销后写入 Cloud Storage，多实例 / 重启后状态仍一致（Cloud Run 必配）
_REDEEM_CODES_GCS_URI = os.getenv("REDEEM_CODES_GCS_URI", "").strip()
_REDEEM_LOG_FILE = _BASE_DIR / "redeem_logs.jsonl"
_USER_SERVICE_FILE = _BASE_DIR / "user_services.json"
_USER_ACCOUNTS_FILE = _BASE_DIR / "user_accounts.json"
_WEBHOOK_RETRY_FILE = _BASE_DIR / "partner_webhook_retry.jsonl"
_TEMP_EXPORT_DIR = Path(tempfile.gettempdir()) / "smart-teacher-assistant-exports"
_TEMP_EXPORT_DIR.mkdir(parents=True, exist_ok=True)
_TEMP_EXPORTS: dict[str, dict[str, str]] = {}
_TEMP_TEMPLATE_DIR = Path(tempfile.gettempdir()) / "smart-teacher-assistant-templates"
_TEMP_TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
_TEMP_TEMPLATES: dict[str, dict[str, str]] = {}
_DEFAULT_APP_STATS = {
    "home_visits": 0,
    "module_clicks": 0,
    "feedback_count": 0,
    "register_count": 0,
}

_DEFAULT_REDEEM_CODES = [
    {
        "code": "VIP2026",
        "status": "unused",
        "token_type": "auto",
        "expires_at": "2026-12-31T23:59:59+00:00",
        "service": {"type": "membership", "name": "会员", "days": 30},
        "description": "30天会员",
    },
    {
        "code": "CZ100",
        "status": "unused",
        "token_type": "auto",
        "expires_at": "2026-12-31T23:59:59+00:00",
        "service": {"type": "balance", "name": "充值", "amount": 100},
        "description": "充值100元",
    },
    {
        "code": "TIMES20",
        "status": "unused",
        "token_type": "auto",
        "expires_at": "2026-12-31T23:59:59+00:00",
        "service": {"type": "quota", "name": "次数", "amount": 20},
        "description": "增加20次",
    },
]


def _parse_partner_tokens(raw: str) -> set[str]:
    items = [part.strip() for part in str(raw or "").split(",")]
    return {item for item in items if item}


PARTNER_REDEEM_TOKENS = _parse_partner_tokens(os.getenv("PARTNER_REDEEM_TOKENS", ""))
PARTNER_REDEEM_SOURCE = str(os.getenv("PARTNER_REDEEM_SOURCE", "third_party_mall")).strip() or "third_party_mall"


def _parse_webhook_urls(raw: str) -> dict[str, str]:
    """解析 PARTNER_WEBHOOK_URLS 环境变量，格式：token1:url1,token2:url2"""
    result: dict[str, str] = {}
    for part in str(raw or "").split(","):
        part = part.strip()
        if not part:
            continue
        token, sep, url = part.partition(":")
        if sep and token.strip() and url.strip():
            result[token.strip()] = url.strip()
    return result


PARTNER_WEBHOOK_URLS: dict[str, str] = _parse_webhook_urls(os.getenv("PARTNER_WEBHOOK_URLS", ""))

TEMPLATE_STANDARD_V116 = {
    "version": "v1.2.0",
    "principles": [
        "星期表头行动态识别（含≥3个星期列），非整表首行硬保护；首列默认保护",
        "上色单元格视为固定模板区，不可覆盖",
        "仅对可填槽位进行净空和内容写入",
        "保留字体、字号、段落、边框、合并关系等样式",
    ],
    "recognition_order": [
        "文档类型识别（周计划/日教案/活动计划）",
        "主标签识别（主题/目标/准备/过程/反思）",
        "主行主列到子槽位映射（含拆分单元格）",
        "内容写入与净空（仅 fillable slots）",
    ],
    "keyword_groups": {
        "meta": ["班级", "日期", "天气", "活动主题", "教学主题", "教育理念"],
        "plan": ["活动目标", "活动准备", "学习活动", "区域活动", "户外活动", "评价与反思"],
        "daily": ["活动导入", "活动过程", "活动延伸", "活动反思", "观察要点"],
    },
}


def _load_template_stats() -> dict[str, int]:
    if not _TEMPLATE_STATS_FILE.exists():
        return dict(_DEFAULT_TEMPLATE_STATS)
    try:
        data = json.loads(_TEMPLATE_STATS_FILE.read_text(encoding="utf-8"))
        out = dict(_DEFAULT_TEMPLATE_STATS)
        if isinstance(data, dict):
            for k in out:
                v = data.get(k, 0)
                out[k] = int(v) if isinstance(v, (int, float, str)) else 0
        return out
    except Exception:
        return dict(_DEFAULT_TEMPLATE_STATS)


def _load_app_stats() -> dict[str, int]:
    if not _APP_STATS_FILE.exists():
        return dict(_DEFAULT_APP_STATS)
    try:
        data = json.loads(_APP_STATS_FILE.read_text(encoding="utf-8"))
        out = dict(_DEFAULT_APP_STATS)
        if isinstance(data, dict):
            for k in out:
                out[k] = int(data.get(k, out[k]) or 0)
        return out
    except Exception:
        return dict(_DEFAULT_APP_STATS)


def _save_app_stats(stats: dict[str, int]) -> None:
    try:
        _APP_STATS_FILE.write_text(
            json.dumps(stats, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    except Exception as e:
        logger.warning("保存应用统计失败：%s", e)


def _inc_app_stat(key: str, delta: int = 1) -> dict[str, int]:
    stats = _load_app_stats()
    if key not in stats:
        stats[key] = 0
    stats[key] = max(0, int(stats.get(key, 0)) + int(delta))
    _save_app_stats(stats)
    return stats


def _save_template_stats(stats: dict[str, int]) -> None:
    try:
        _TEMPLATE_STATS_FILE.write_text(
            json.dumps(stats, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    except Exception:
        # 计数写入失败不影响主流程下载
        pass


def _inc_template_download(template_id: str) -> dict[str, int]:
    stats = _load_template_stats()
    if template_id not in stats:
        stats[template_id] = 0
    stats[template_id] += 1
    _save_template_stats(stats)
    return stats


def _append_jsonl(path: Path, payload: dict) -> None:
    try:
        with path.open("a", encoding="utf-8") as f:
            f.write(json.dumps(payload, ensure_ascii=False) + "\n")
    except Exception as e:
        logger.warning("写入 JSONL 失败（%s）：%s", path.name, e)


def _utc_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _load_registered_ids() -> set[str]:
    ids: set[str] = set()
    if not _REGISTER_LOG_FILE.exists():
        return ids
    try:
        for line in _REGISTER_LOG_FILE.read_text(encoding="utf-8").splitlines():
            if not line.strip():
                continue
            try:
                row = json.loads(line)
            except Exception:
                continue
            rid = str(row.get("identifier", "")).strip().lower()
            if rid:
                ids.add(rid)
    except Exception as e:
        logger.warning("读取注册记录失败：%s", e)
    return ids


def _parse_gs_uri(uri: str) -> tuple[str, str]:
    u = str(uri or "").strip()
    if not u.startswith("gs://"):
        raise ValueError("invalid gs uri")
    rest = u[5:]
    if "/" not in rest:
        raise ValueError("invalid gs uri")
    bucket, _, blob = rest.partition("/")
    if not bucket or not blob:
        raise ValueError("invalid gs uri")
    return bucket, blob


def _redeem_codes_dict_from_json_list(raw: object) -> dict[str, dict]:
    items = raw if isinstance(raw, list) else []
    out: dict[str, dict] = {}
    for item in items:
        if isinstance(item, dict) and item.get("code"):
            out[str(item["code"]).strip().upper()] = dict(item)
    return out


def _merge_redeem_row(a: dict, b: dict) -> dict:
    """合并两条卡密：任一侧为已使用则保留已使用，避免并发核销互相覆盖。"""
    sa = str(a.get("status", "")).strip().lower()
    sb = str(b.get("status", "")).strip().lower()
    if sa == "used" and sb != "used":
        return dict(a)
    if sb == "used" and sa != "used":
        return dict(b)
    if sa == "used" and sb == "used":
        ta = str(a.get("used_at_utc", ""))
        tb = str(b.get("used_at_utc", ""))
        return dict(a) if ta >= tb else dict(b)
    return dict(b)


def _merge_redeem_dicts(remote: dict[str, dict], local: dict[str, dict]) -> dict[str, dict]:
    keys = set(remote) | set(local)
    out: dict[str, dict] = {}
    for k in keys:
        if k in remote and k in local:
            out[k] = _merge_redeem_row(remote[k], local[k])
        elif k in remote:
            out[k] = dict(remote[k])
        else:
            out[k] = dict(local[k])
    return out


def _try_load_redeem_codes_gcs() -> dict[str, dict] | None:
    """从 GCS 读取卡密库；对象不存在返回 None（由调用方回退本地）。"""
    if not _REDEEM_CODES_GCS_URI:
        return None
    try:
        from google.cloud import storage  # noqa: PLC0415
    except ImportError:
        logger.warning("已设置 REDEEM_CODES_GCS_URI 但未安装 google-cloud-storage，跳过 GCS")
        return None
    try:
        bucket_name, blob_name = _parse_gs_uri(_REDEEM_CODES_GCS_URI)
        client = storage.Client()
        bucket = client.bucket(bucket_name)
        blob = bucket.blob(blob_name)
        if not blob.exists():
            return None
        raw = json.loads(blob.download_as_text(encoding="utf-8"))
        return _redeem_codes_dict_from_json_list(raw)
    except Exception as e:
        logger.warning("从 GCS 读取卡密库失败：%s", e)
        return None


def _save_redeem_codes_gcs_merged(incoming: dict[str, dict]) -> dict[str, dict] | None:
    """将 incoming 与 GCS 当前内容合并后写回；带乐观锁重试。成功返回合并后的全量 dict。"""
    try:
        from google.api_core import exceptions as gexc  # noqa: PLC0415
        from google.cloud import storage  # noqa: PLC0415
    except ImportError:
        logger.warning("无法写入 GCS：缺少 google-cloud-storage")
        return None
    bucket_name, blob_name = _parse_gs_uri(_REDEEM_CODES_GCS_URI)
    client = storage.Client()
    bucket = client.bucket(bucket_name)
    blob = bucket.blob(blob_name)
    for attempt in range(16):
        try:
            remote: dict[str, dict] = {}
            gen: int | None = None
            if blob.exists():
                blob.reload()
                gen = blob.generation
                raw = json.loads(blob.download_as_text(encoding="utf-8"))
                remote = _redeem_codes_dict_from_json_list(raw)
            merged = _merge_redeem_dicts(remote, incoming)
            payload = json.dumps(list(merged.values()), ensure_ascii=False, indent=2)
            if gen is not None:
                blob.upload_from_string(
                    payload,
                    content_type="application/json; charset=utf-8",
                    if_generation_match=gen,
                )
            else:
                blob.upload_from_string(payload, content_type="application/json; charset=utf-8")
            return merged
        except gexc.PreconditionFailed:
            continue
        except Exception as e:
            logger.warning("写入 GCS 卡密库失败（第 %s 次）：%s", attempt + 1, e)
            if attempt >= 15:
                return None
    return None


def _load_redeem_codes() -> dict[str, dict]:
    if _REDEEM_CODES_GCS_URI:
        gcs_data = _try_load_redeem_codes_gcs()
        if gcs_data is not None:
            return gcs_data
        logger.info("GCS 卡密库不存在或暂不可读，回退本地 redeem_codes.json")

    if not _REDEEM_CODES_FILE.exists():
        try:
            _REDEEM_CODES_FILE.write_text(
                json.dumps(_DEFAULT_REDEEM_CODES, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
        except Exception as e:
            logger.warning("初始化卡密库失败：%s", e)
            return {item["code"]: dict(item) for item in _DEFAULT_REDEEM_CODES}

    try:
        raw = json.loads(_REDEEM_CODES_FILE.read_text(encoding="utf-8"))
        return _redeem_codes_dict_from_json_list(raw)
    except Exception as e:
        logger.warning("读取卡密库失败：%s", e)
        return {item["code"]: dict(item) for item in _DEFAULT_REDEEM_CODES}


def _save_redeem_codes(codes: dict[str, dict]) -> None:
    to_write: dict[str, dict] = dict(codes)
    if _REDEEM_CODES_GCS_URI:
        merged = _save_redeem_codes_gcs_merged(codes)
        if merged is not None:
            to_write = merged
        else:
            logger.warning("GCS 卡密库写入失败，仅写本地副本（多实例下可能不一致）")
    try:
        _REDEEM_CODES_FILE.write_text(
            json.dumps(list(to_write.values()), ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    except Exception as e:
        logger.warning("保存卡密库失败：%s", e)


def _load_user_services() -> dict[str, dict]:
    """读用户服务：优先 Firestore，回退本地 JSON。"""
    if FIRESTORE_ENABLED:
        try:
            docs = _fs().collection("user_services").stream()
            result = {}
            for doc in docs:
                result[doc.id] = doc.to_dict()
            if result:
                return result
        except Exception as e:
            logger.warning("Firestore 读取用户服务失败，回退本地：%s", e)
    if not _USER_SERVICE_FILE.exists():
        return {}
    try:
        data = json.loads(_USER_SERVICE_FILE.read_text(encoding="utf-8"))
        return data if isinstance(data, dict) else {}
    except Exception as e:
        logger.warning("读取用户服务失败：%s", e)
        return {}


def _save_user_services(data: dict[str, dict]) -> None:
    """写用户服务：同时写 Firestore 和本地 JSON（双保险）。"""
    # 写 Firestore
    if FIRESTORE_ENABLED:
        try:
            batch = _fs().batch()
            col = _fs().collection("user_services")
            for user_id, entry in data.items():
                batch.set(col.document(user_id), entry)
            batch.commit()
        except Exception as e:
            logger.warning("Firestore 写用户服务失败：%s", e)
    # 写本地 JSON 兜底
    try:
        _USER_SERVICE_FILE.write_text(
            json.dumps(data, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    except Exception as e:
        logger.warning("保存用户服务失败：%s", e)


# ──────────────────────────────────────────────
# 卡密自动生成
# ──────────────────────────────────────────────
_CODE_ALPHABET = "".join(c for c in (string.ascii_uppercase + string.digits) if c not in "OI15S")
_CODE_PREFIX_MAP = {"membership": "M", "balance": "B", "quota": "Q"}


def _generate_code(service_type: str = "", length: int = 8) -> str:
    prefix = _CODE_PREFIX_MAP.get(str(service_type).strip(), "X")
    body = "".join(secrets.choice(_CODE_ALPHABET) for _ in range(length))
    return f"{prefix}{body}"


def _generate_unique_code(codes: dict, service_type: str = "", length: int = 8, max_retries: int = 20) -> str:
    for _ in range(max_retries):
        code = _generate_code(service_type, length)
        if code not in codes:
            return code
    raise RuntimeError("卡密生成碰撞次数过多，请检查库存量")


# ──────────────────────────────────────────────
# Webhook 回调 & 重试队列
# ──────────────────────────────────────────────

def _load_webhook_retries() -> list[dict]:
    if not _WEBHOOK_RETRY_FILE.exists():
        return []
    items: list[dict] = []
    try:
        for line in _WEBHOOK_RETRY_FILE.read_text(encoding="utf-8").splitlines():
            line = line.strip()
            if line:
                try:
                    items.append(json.loads(line))
                except Exception:
                    pass
    except Exception as e:
        logger.warning("读取 Webhook 重试队列失败：%s", e)
    return items


def _save_webhook_retries(items: list[dict]) -> None:
    try:
        _WEBHOOK_RETRY_FILE.write_text(
            "\n".join(json.dumps(item, ensure_ascii=False) for item in items),
            encoding="utf-8",
        )
    except Exception as e:
        logger.warning("保存 Webhook 重试队列失败：%s", e)


def _enqueue_webhook_retry(payload: dict) -> None:
    items = _load_webhook_retries()
    items.append(payload)
    _save_webhook_retries(items)


async def _fire_webhook_once(url: str, body: dict) -> bool:
    """向第三方发送一次 Webhook POST，返回是否成功（在线程池中执行，不阻塞事件循环）。"""
    import urllib.request
    import urllib.error

    data = json.dumps(body, ensure_ascii=False).encode("utf-8")

    def _do_post() -> bool:
        req = urllib.request.Request(
            url,
            data=data,
            headers={
                "Content-Type": "application/json",
                "User-Agent": "SmartTeacher-Webhook/1.0",
            },
            method="POST",
        )
        try:
            with urllib.request.urlopen(req, timeout=10) as resp:
                return resp.status < 300
        except Exception as exc:
            logger.warning("Webhook POST 失败 url=%s: %s", url, exc)
            return False

    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _do_post)


async def _dispatch_webhook(callback_url: str, body: dict, order_id: str, code: str) -> None:
    """核销后异步触发 Webhook；失败则写入重试队列。"""
    if not callback_url:
        return
    success = await _fire_webhook_once(callback_url, body)
    if not success:
        retry_item = {
            "id": f"WH-{uuid4().hex[:10]}",
            "callback_url": callback_url,
            "body": body,
            "order_id": order_id,
            "code": code,
            "retry_count": 0,
            "last_attempt_utc": _utc_iso(),
            "status": "pending",
            "created_at_utc": _utc_iso(),
        }
        _enqueue_webhook_retry(retry_item)
        logger.warning("Webhook 首次失败，已加入重试队列 code=%s order=%s", code, order_id)


async def _webhook_retry_loop() -> None:
    """后台定时任务：每 2 分钟重试一次失败的 Webhook，累计 3 次失败后标记 error 并停止重试。"""
    while True:
        await asyncio.sleep(120)
        items = _load_webhook_retries()
        if not items:
            continue
        remaining: list[dict] = []
        for item in items:
            if item.get("status") == "error":
                remaining.append(item)
                continue
            retry_count = int(item.get("retry_count", 0))
            if retry_count >= 3:
                item["status"] = "error"
                logger.error(
                    "Webhook 三次失败已放弃 code=%s order=%s",
                    item.get("code"), item.get("order_id"),
                )
                remaining.append(item)
                continue
            success = await _fire_webhook_once(item.get("callback_url", ""), item.get("body", {}))
            item["retry_count"] = retry_count + 1
            item["last_attempt_utc"] = _utc_iso()
            if success:
                logger.info(
                    "Webhook 重试成功 code=%s order=%s 第%d次",
                    item.get("code"), item.get("order_id"), item["retry_count"],
                )
            else:
                remaining.append(item)
        _save_webhook_retries([i for i in remaining if i.get("status") in ("pending", "error")])


# ──────────────────────────────────────────────
# 用户账户（极简内测版：手机号即 user_id）
# ──────────────────────────────────────────────

def _load_user_accounts() -> dict[str, dict]:
    """读用户账户：优先 Firestore，回退本地 JSON。"""
    if FIRESTORE_ENABLED:
        try:
            docs = _fs().collection("users").stream()
            result = {}
            for doc in docs:
                result[doc.id] = doc.to_dict()
            if result:
                return result
        except Exception as e:
            logger.warning("Firestore 读取用户账户失败，回退本地：%s", e)
    if not _USER_ACCOUNTS_FILE.exists():
        return {}
    try:
        data = json.loads(_USER_ACCOUNTS_FILE.read_text(encoding="utf-8"))
        return data if isinstance(data, dict) else {}
    except Exception as e:
        logger.warning("读取用户账户失败：%s", e)
        return {}


def _save_user_accounts(data: dict[str, dict]) -> None:
    """写用户账户：同时写 Firestore 和本地 JSON（双保险）。"""
    if FIRESTORE_ENABLED:
        try:
            batch = _fs().batch()
            col = _fs().collection("users")
            for uid, entry in data.items():
                batch.set(col.document(uid), entry, merge=True)
            batch.commit()
        except Exception as e:
            logger.warning("Firestore 写用户账户失败：%s", e)
    try:
        _USER_ACCOUNTS_FILE.write_text(
            json.dumps(data, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    except Exception as e:
        logger.warning("保存用户账户失败：%s", e)


def _get_or_create_user(openid: str) -> dict:
    """获取或创建用户档案（含 agent_profile）。"""
    accounts = _load_user_accounts()
    if openid not in accounts:
        accounts[openid] = {
            "openid": openid,
            "user_id": None,
            "created_at": _utc_iso(),
            "agent_profile": {
                "name": "小助手",
                "personality": "热心、耐心",
                "tone": "亲切温暖",
                "style": "鼓励式教学",
            },
            "last_token": None,
        }
        _save_user_accounts(accounts)
    return accounts[openid]


def _generate_user_token(openid: str) -> str:
    """生成用户会话 token。"""
    return f"ut_{uuid4().hex}"


def _verify_user_token(token: str) -> str:
    """验证 user_token，返回对应 openid。失败抛 401。"""
    token = str(token or "").strip()
    if not token:
        raise HTTPException(status_code=401, detail="请先登录（user_token 缺失）")
    if FIRESTORE_ENABLED:
        try:
            docs = _fs().collection("users").where("last_token", "==", token).limit(1).stream()
            for doc in docs:
                return doc.id
        except Exception as e:
            logger.warning("Firestore token 验证失败，回退本地：%s", e)
    # 本地回退
    accounts = _load_user_accounts()
    for openid, entry in accounts.items():
        if entry.get("last_token") == token:
            return openid
    raise HTTPException(status_code=401, detail="token 无效或已过期，请重新登录")


def _append_redeem_log(payload: dict) -> None:
    _append_jsonl(_REDEEM_LOG_FILE, payload)


def _read_json_file(path: Path) -> dict:
    if not path.exists():
        return {}
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
        return payload if isinstance(payload, dict) else {}
    except Exception:
        return {}


def _knowledge_base_status() -> dict:
    index_payload = _read_json_file(_KNOWLEDGE_INDEX_FILE)
    route_payload = _read_json_file(_KNOWLEDGE_ROUTE_FILE)
    records = index_payload.get("records", [])
    bucket_counts = index_payload.get("buckets", {})
    source_counts: dict[str, int] = {}
    if isinstance(records, list):
        for item in records:
            if not isinstance(item, dict):
                continue
            tier = str(item.get("source_tier", "unknown") or "unknown")
            source_counts[tier] = source_counts.get(tier, 0) + 1

    routes = route_payload.get("profile_routes", {})
    profile_count = len(routes) if isinstance(routes, dict) else 0
    generated_at = (
        str(index_payload.get("generated_at_utc", "")).strip()
        or str(route_payload.get("generated_at_utc", "")).strip()
    )
    return {
        "enabled": _KNOWLEDGE_BASE_DIR.exists(),
        "doc_count": int(index_payload.get("doc_count", 0) or 0),
        "bucket_counts": bucket_counts if isinstance(bucket_counts, dict) else {},
        "source_counts": source_counts,
        "profile_route_count": profile_count,
        "generated_at_utc": generated_at,
        "index_ready": bool(index_payload),
    }


def _redeem_code_core(raw_code: str, user_id: str, source: str, extra_log: Optional[dict] = None) -> dict:
    code = str(raw_code or "").strip().upper()
    account = str(user_id or "").strip().lower()
    redeem_source = str(source or "").strip() or "unknown"
    if not code:
        raise HTTPException(status_code=400, detail="请填写卡密")
    if not account:
        raise HTTPException(status_code=400, detail="请填写账号或手机号")

    codes = _load_redeem_codes()
    item = codes.get(code)
    log_payload = {
        "id": f"RD-{uuid4().hex[:10]}",
        "created_at_utc": _utc_iso(),
        "code": code,
        "user_id": account,
        "source": redeem_source,
    }
    if isinstance(extra_log, dict):
        log_payload.update(extra_log)

    if not item:
        log_payload["result"] = "invalid"
        _append_redeem_log(log_payload)
        return {"ok": False, "status": "invalid", "message": "无效"}

    expires_at = str(item.get("expires_at", "")).strip()
    expire_dt = None
    try:
        expire_dt = datetime.fromisoformat(expires_at)
    except Exception:
        pass

    if item.get("status") == "used":
        log_payload["result"] = "used"
        _append_redeem_log(log_payload)
        return {"ok": False, "status": "used", "message": "已使用"}

    if expire_dt is not None and datetime.now(timezone.utc) > expire_dt:
        log_payload["result"] = "expired"
        _append_redeem_log(log_payload)
        return {"ok": False, "status": "expired", "message": "已过期"}

    service = item.get("service", {})
    user_services = _load_user_services()
    entry = user_services.setdefault(
        account,
        {"membership_until": None, "balance": 0, "quota": 0, "rewards": []},
    )

    granted: dict[str, object] = {
        "type": service.get("type", ""),
        "name": service.get("name", ""),
        "code": code,
        "granted_at_utc": _utc_iso(),
    }
    service_type = str(service.get("type", "")).strip()
    now = datetime.now(timezone.utc)
    if service_type == "membership":
        days = int(service.get("days", 0) or 0)
        until = now
        if entry.get("membership_until"):
            try:
                prev = datetime.fromisoformat(str(entry["membership_until"]))
                if prev > now:
                    until = prev
            except Exception:
                pass
        until = until + timedelta(days=days)
        entry["membership_until"] = until.isoformat()
        granted["membership_until"] = until.isoformat()
    elif service_type == "balance":
        amount = int(service.get("amount", 0) or 0)
        entry["balance"] = int(entry.get("balance", 0) or 0) + amount
        granted["balance_added"] = amount
        granted["balance_total"] = entry["balance"]
    elif service_type == "quota":
        amount = int(service.get("amount", 0) or 0)
        entry["quota"] = int(entry.get("quota", 0) or 0) + amount
        granted["quota_added"] = amount
        granted["quota_total"] = entry["quota"]
    entry.setdefault("rewards", []).append(granted)
    user_services[account] = entry
    _save_user_services(user_services)

    item["status"] = "used"
    item["used_at_utc"] = _utc_iso()
    item["used_by"] = account
    item["used_source"] = redeem_source
    codes[code] = item
    _save_redeem_codes(codes)

    log_payload["result"] = "success"
    log_payload["service"] = service
    _append_redeem_log(log_payload)
    result: dict = {
        "ok": True,
        "status": "success",
        "message": "成功",
        "service": service,
        "granted": granted,
    }
    if item.get("callback_url"):
        result["_callback_url"] = str(item["callback_url"])
        result["_order_id"] = str(item.get("order_id", ""))
    return result


def _build_mini_doc_payload(
    filled_bytes: bytes,
    original_name: str,
    export_engine: str,
) -> dict[str, str]:
    token = uuid4().hex
    export_path = _TEMP_EXPORT_DIR / f"{token}.docx"
    export_path.write_bytes(filled_bytes)
    _TEMP_EXPORTS[token] = {
        "path": str(export_path),
        "filename": original_name,
        "engine": export_engine,
    }
    return {
        "status": "ok",
        "download_url": f"/mini-export/{token}",
        "filename": original_name,
        "engine": export_engine,
        "file_base64": base64.b64encode(filled_bytes).decode("ascii"),
    }

# ──────────────────────────────────────────────
# 关键字映射表：单元格文本 → 填充字段 key
# 优先级从上到下，第一个匹配的规则生效
# ──────────────────────────────────────────────
CELL_KEYWORD_MAP: list[tuple[list[str], str]] = [
    # (关键词列表,          字段 key)
    (["基础信息", "基本信息"],                                     "class_info"),
    (["上周情况分析", "上周分析", "上周情况"],                      "children_baseline"),
    (["本周重点与难点", "重点与难点", "重点难点"],                  "key_difficulty"),
    (["幼儿已有经验", "已有经验", "前测"],                          "children_baseline"),
    (["本周活动总览表", "周活动总览", "周安排总览"],                "week_overview"),
    (["每日活动要点", "每日要点"],                                  "daily_points"),
    (["户外与体能活动", "体能活动"],                                "outdoor"),
    (["生活活动与保育", "保育"],                                    "life"),
    (["环境创设", "环创活动", "环境布置", "环创"],                  "environment"),
    (["家园共育", "家园互动", "家长工作", "家园社区互动", "家园活动", "家园"], "family"),
    (["个别化支持", "个别指导"],                                    "individual_support"),
    (["安全与风险提示", "安全提示", "风险提示", "安全"],            "safety_risk"),
    (["资源与材料清单", "材料清单", "资源清单"],                    "resource_list"),
    (["观察记录计划", "观察计划"],                                  "observation_plan"),
    (["周反思", "本周反思"],                                        "evaluation"),
    (["下周衔接", "下周计划", "衔接"],                              "next_week_plan"),
    (["午睡指导", "睡姿指导", "穿脱衣物", "衣物摆放", "睡前习惯"],  "nap_guidance"),
    (["教学主题", "活动主题", "本月主题", "单元主题", "主题"],   "theme"),
    (["教育理念", "课程理念", "理念", "风格"],                   "philosophy"),
    (["活动目标", "教学目标", "重点目标", "目标"],               "goals"),
    (["活动准备", "材料准备", "准备"],                           "preparation"),
    (["指导要点"],                                               "guidance"),
    (["晨间运动", "晨间", "早谈", "早  谈", "晨谈", "早操"],      "morning"),
    (["户外活动", "体育活动", "户外游戏", "户外"],               "outdoor"),
    (["环创活动", "环境创设", "环创"],                           "environment"),
    (["生活活动", "一日生活", "生活"],                           "life"),
    (["学习活动", "集中活动", "教学活动"],                       "study"),
    (["游戏活动", "游戏"],                                       "game"),
    (["区域活动", "区角活动", "区域"],                           "area"),
    (["离园活动", "离园"],                                       "departure"),
    (["幼儿自主", "自主发起"],                                   "child_initiative"),
    (["评价与反思", "评价", "反思", "小结"],                     "evaluation"),
    (["班级", "班"],                                             "class_info"),
]

# 可填到周网格内容格的活动字段集合（不含 meta 字段）
_ACTIVITY_FIELDS = frozenset({
    "morning", "outdoor", "environment", "life", "study",
    "game", "area", "family", "departure", "evaluation", "nap_guidance",
})

WEEKLY_STANDARD_MODULES = [
    "基础信息",
    "本周主题",
    "周总目标（五大领域）",
    "本周重点与难点",
    "幼儿已有经验",
    "本周活动总览表",
    "每日活动要点",
    "区域活动设计",
    "户外与体能活动",
    "生活活动与保育",
    "环境创设",
    "家园共育",
    "个别化支持",
    "安全与风险提示",
    "资源与材料清单",
    "观察记录计划",
    "周反思",
    "下周衔接",
    "午睡指导",
]


def match_field(cell_text: str) -> Optional[str]:
    """根据单元格文本匹配应填充的字段 key，无匹配返回 None。
    同时检查原文和去空行拼合版，兼容 Word 多段拆字（如"集体\\n\\n活动"）。"""
    t = cell_text.strip()
    t_compact = "".join(t.split())   # 去掉所有空白拼合，用于匹配拆行关键词
    for keywords, field in CELL_KEYWORD_MAP:
        for kw in keywords:
            if kw in t or kw in t_compact:
                return field
    return None


# ──────────────────────────────────────────────
# 教育理念专业词汇库（补充 Prompt 提示词）
# ──────────────────────────────────────────────
PHILOSOPHY_HINTS: dict[str, str] = {
    "蒙氏教育（AMI/AMS）": (
        "请大量使用以下专业术语：操作教具、敏感期观察、三段式教学、工作周期、"
        "有准备的环境（Prepared Environment）、工作毯、混龄协作、内在纪律感。"
    ),
    "瑞吉欧教育": (
        "请大量使用以下专业术语：生成课程（Emergent Curriculum）、环境留痕（Documentation）、"
        "一百种语言、项目网络（Project Web）、呈现板（Documentation Panel）、协作解读。"
    ),
    "DAP 发展适宜性实践": (
        "请大量使用以下专业术语：发展适宜性、年龄适宜性、个体适宜性、"
        "最近发展区（ZPD）、支架式学习（Scaffolding）、真实性评估、文化回应性教学。"
    ),
    "华德福教育": (
        "请大量使用以下专业术语：生命节律、季节庆典、季节桌（Seasonal Table）、"
        "意志力、优律思美（Eurythmy）、故事讲述（Storytelling）、吸气-呼气节奏。"
    ),
    "项目化学习（PBL）": (
        "请大量使用以下专业术语：驱动性问题（Driving Question）、成果展示、"
        "评价量规（Rubric）、跨领域整合、真实受众、合作探究。"
    ),
    "自主游戏 / 游戏化课程": (
        "请大量使用以下专业术语：儿童视角、游戏观察、松散材料（Loose Parts）、"
        "低结构材料、游戏即工作、自由探索。"
    ),
    "传统文化 / 国学教育": (
        "请大量使用以下专业术语：二十四节气、节气文化、经典诵读、传统礼仪、"
        "传统游戏传承、非遗体验、文化认同。"
    ),
    "五大领域": (
        "请对应五大领域（健康、语言、社会、科学、艺术）分别阐述核心经验，"
        "并参照《3-6岁儿童学习与发展指南》的典型表现。"
    ),
}


# ──────────────────────────────────────────────
# AI 内容生成
# ──────────────────────────────────────────────
def build_system_prompt() -> str:
    return textwrap.dedent("""
        你是一位拥有10年经验的资深幼儿园教研主任，擅长编写专业的幼儿园活动计划。
        你的输出必须：
        1. 严格按照 JSON 格式返回，不包含任何 Markdown 代码块标记
        2. 内容专业、温暖、具有可操作性
        3. 每条目标/环节内容言简意赅，控制在 80 字以内
        4. 评价部分要有具体的观察维度和记录建议
        5. 必须紧扣“教学主题”和“教育理念”，禁止输出与主题无关的泛化套话
        6. 不要编造模板中不存在的额外模块，只填充给定 JSON 结构
    """).strip()


def build_user_prompt(
    theme: str,
    phil: str,
    activities: list[str],
    child_initiative: bool,
    child_desc: str,
    template_outline: Optional[list[str]] = None,
    class_level: str = "",
) -> str:
    phil_hint = PHILOSOPHY_HINTS.get(phil, "")
    class_hint = CLASS_LEVEL_HINTS.get(class_level, "")
    acts_str = "、".join(activities) if activities else "区域活动、户外活动"
    outline_hint = ""
    if template_outline:
        outline_items = "\n".join(f"- {x}" for x in template_outline[:14])
        outline_hint = f"\n【老师模板提纲参考（优先对齐语义结构，不要求复刻版式）】\n{outline_items}\n"

    return textwrap.dedent(f"""
        请为以下幼儿园活动生成专业内容，以 JSON 格式返回。

        【输入信息】
        - 教学主题：{theme}
        - 教育理念 / 园本特色：{phil}
        - 班级：{class_level or "未指定（请按中班水平默认）"}
        - 活动重点：{acts_str}
        - 幼儿自主发起活动：{"是，" + child_desc if child_initiative and child_desc else "否"}

        【理念专业词汇要求】
        {phil_hint}

        【班级年龄特征约束】
        {class_hint or "按中班（4-5岁）水平默认生成，活动难度与目标表述取中间值。"}
        {outline_hint}
        【贴合要求（严格）】
        - 所有字段必须围绕「{theme}」，不能写成任意主题都通用的话术。
        - 优先贴合老师模板提纲语义；若提纲含“周一~周五/观察/反思”等，应在对应字段体现。
        - 周一到周五应体现五大领域（健康、语言、社会、科学、艺术）的均衡分配，不可集中在单一领域。
        - 输出应体现园本特色，不要写成空泛口号。
        - 仅返回可被 json.loads 解析的对象 JSON。

        【JSON 输出格式（严格遵守）】
        {{
          "weekly_targets": {{
            "teaching": "本周教学目标（贴合《3-6岁指南》并结合班级现状）",
            "life": "本周生活目标（卫生保健、自理、安全习惯）",
            "family": "本周家园共育目标（家庭配合与沟通）",
            "environment": "本周环创目标（为主题学习提供环境与材料支持）"
          }},
          "goals": [
            "目标1（含领域/维度前缀）",
            "目标2",
            "目标3",
            "目标4（可选）",
            "目标5（可选）"
          ],
          "preparation": [
            "材料1",
            "材料2",
            "材料3"
          ],
          "activities": {{
            "morning":     "晨间运动的具体教学建议（如未选择则留空字符串）",
            "outdoor":     "户外活动的具体教学建议",
            "environment": "环创活动的具体教学建议",
            "life":        "生活活动的具体教学建议",
            "study":       "集中活动每天的活动名称。【班级规则】小班（≤4岁）只有1个；中班/大班各有2个（须动静结合，如语言+体育）。格式：周一 领域1：《活动名1》；领域2：《活动名2》（仅中班/大班）\\n周二……\\n……（共五天，只写名称不写目标）",
            "game":        "游戏活动每天的游戏名称，格式：周一 《游戏名》\\n周二 《游戏名》\\n……（只写名称，共五天）",
            "area":        "区域活动的具体教学建议",
            "family":      "家园活动的具体教学建议",
            "departure":   "离园活动的具体教学建议"
          }},
          "child_initiative_note": "如果幼儿有自主发起活动，给出教师跟进建议（无则留空）",
          "nap_guidance": "午睡指导（睡姿、穿脱衣物、衣物摆放、睡前习惯）",
          "evaluation": "评价与反思建议（含观察维度、记录方式、改进方向）"
        }}

        只返回 JSON，不要任何其他文字。
    """).strip()


_OUTLINE_KEYS = (
    "主题", "目标", "准备", "过程", "导入", "延伸", "反思", "观察",
    "评价", "家园", "晨间", "户外", "区域", "离园", "周一", "周二", "周三", "周四", "周五",
)


def _extract_template_outline(template_bytes: bytes, max_items: int = 14) -> list[str]:
    """
    仅抽取“内容提纲语义”，不依赖版式。用于让生成内容与老师模板语义更一致。
    """
    try:
        doc = Document(io.BytesIO(template_bytes))
    except Exception:
        return []

    def _clean(s: str) -> str:
        t = re.sub(r"\s+", " ", (s or "")).strip()
        return t.strip("：:;；,.。")

    rows: list[str] = []
    seen: set[str] = set()

    def push(raw: str) -> None:
        t = _clean(raw)
        if not t or t in seen:
            return
        if len(t) < 2 or len(t) > 26:
            return
        if not any(k in t for k in _OUTLINE_KEYS):
            return
        seen.add(t)
        rows.append(t)

    for p in doc.paragraphs:
        push(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                push(cell.text)
    return rows[:max_items]


_ACTIVITY_KEYS = (
    "morning", "outdoor", "environment", "life", "study", "game", "area", "family", "departure"
)


def _normalize_content_payload(
    payload: dict,
    *,
    theme: str,
    phil: str,
    activities: list[str],
    child_initiative: bool,
    child_desc: str,
) -> dict:
    """
    强约束输出：即使模型波动，也保证结构齐全、字段稳定、语义一致。
    """
    fallback = _mock_content(theme, phil, activities)
    data = payload if isinstance(payload, dict) else {}
    weekly_targets_in = data.get("weekly_targets", {})
    weekly_targets_in = weekly_targets_in if isinstance(weekly_targets_in, dict) else {}
    weekly_targets = {
        "teaching": str(weekly_targets_in.get("teaching", "") or "").strip(),
        "life": str(weekly_targets_in.get("life", "") or "").strip(),
        "family": str(weekly_targets_in.get("family", "") or "").strip(),
        "environment": str(weekly_targets_in.get("environment", "") or "").strip(),
    }

    goals_raw = data.get("goals", [])
    goals = [str(x).strip() for x in (goals_raw or []) if str(x).strip()]
    if len(goals) < 3:
        goals = [str(x).strip() for x in fallback["goals"]]
    goals = goals[:5]
    if not weekly_targets["teaching"]:
        weekly_targets["teaching"] = f"围绕「{theme}」落实五大领域核心经验，体现{phil}特色并兼顾班级个体差异。"
    if not weekly_targets["life"]:
        weekly_targets["life"] = "聚焦卫生保健、安全规则与生活自理，形成可持续的日常习惯。"
    if not weekly_targets["family"]:
        weekly_targets["family"] = "通过家园沟通与家庭延伸任务，形成教育一致性与共同支持。"
    if not weekly_targets["environment"]:
        weekly_targets["environment"] = f"基于「{theme}」优化主题角与材料投放，支持角色扮演与探究活动。"

    prep_raw = data.get("preparation", [])
    preparation = [str(x).strip() for x in (prep_raw or []) if str(x).strip()]
    if len(preparation) < 3:
        preparation = [str(x).strip() for x in fallback["preparation"]]
    preparation = preparation[:5]

    in_acts = data.get("activities", {})
    in_acts = in_acts if isinstance(in_acts, dict) else {}
    selected = set(activities or [])
    normalized_acts: dict[str, str] = {}
    for key in _ACTIVITY_KEYS:
        v = str(in_acts.get(key, "") or "").strip()
        if not v:
            label = ACTIVITY_LABEL_MAP.get(key, key)
            if key in selected:
                v = f"围绕「{theme}」开展{label}，按{phil}理念设计可执行步骤，并记录幼儿关键表现。"
            else:
                v = f"结合「{theme}」预留{label}联动建议，支持后续按班级实际灵活启用。"
        normalized_acts[key] = v

    child_note = str(data.get("child_initiative_note", "") or "").strip()
    if child_initiative:
        if not child_note:
            child_note = (
                f"围绕幼儿自主发起内容进行追问与延展：{child_desc}"
                if child_desc else
                "观察并承接幼儿自主发起线索，采用小步追问与材料支持，形成可延续的探究路径。"
            )
    else:
        child_note = ""

    evaluation = str(data.get("evaluation", "") or "").strip()
    if not evaluation:
        evaluation = str(fallback["evaluation"])
    nap_guidance = str(data.get("nap_guidance", "") or "").strip()
    if not nap_guidance:
        nap_guidance = str(fallback.get("nap_guidance", ""))

    return {
        "weekly_targets": weekly_targets,
        "goals": goals,
        "preparation": preparation,
        "activities": normalized_acts,
        "child_initiative_note": child_note,
        "nap_guidance": nap_guidance,
        "evaluation": evaluation,
    }


def _parse_json_payload(raw_text: str) -> dict:
    """
    兼容模型偶发输出前后缀说明/代码块的情况，尽量抽取首个 JSON 对象。
    """
    raw = str(raw_text or "").strip()
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    try:
        parsed = json.loads(raw)
        return parsed if isinstance(parsed, dict) else {}
    except Exception:
        pass

    start = raw.find("{")
    end = raw.rfind("}")
    if start >= 0 and end > start:
        try:
            parsed = json.loads(raw[start:end + 1])
            return parsed if isinstance(parsed, dict) else {}
        except Exception:
            return {}
    return {}


def generate_content(
    theme: str,
    phil: str,
    activities: list[str],
    child_initiative: bool,
    child_desc: str,
    template_outline: Optional[list[str]] = None,
    class_level: str = "",
) -> dict:
    """调用阿里云百炼 Qwen-Max，返回解析后的内容字典"""
    if not DASHSCOPE_API_KEY:
        if not ALLOW_MOCK_CONTENT:
            raise HTTPException(
                status_code=503,
                detail="AI 服务未配置（缺少 DASHSCOPE_API_KEY）；为保证生成质量，已禁止回退 Mock 内容。",
            )
        # 仅在显式允许时回退 Mock，方便本地开发调试
        logger.warning("AI Key 缺失，ALLOW_MOCK_CONTENT=1，回退 Mock 生成")
        return _normalize_content_payload(
            _mock_content(theme, phil, activities),
            theme=theme,
            phil=phil,
            activities=activities,
            child_initiative=child_initiative,
            child_desc=child_desc,
        )

    messages = [
        {"role": "system", "content": build_system_prompt()},
        {
            "role": "user",
            "content": build_user_prompt(
                theme, phil, activities, child_initiative, child_desc, template_outline,
                class_level=class_level,
            ),
        },
    ]

    try:
        response = client.chat.completions.create(
            model=AI_MODEL,
            messages=messages,
            temperature=0.25,
            max_tokens=4096,
        )
        raw = response.choices[0].message.content.strip()
        parsed = _parse_json_payload(raw)
        if not parsed:
            # 一次轻量重试，强制仅返回 JSON 对象
            retry_messages = messages + [
                {"role": "assistant", "content": raw},
                {
                    "role": "user",
                    "content": "你上一条输出不是可解析 JSON。请仅返回一个合法 JSON 对象，不要附加任何文字。",
                },
            ]
            retry_resp = client.chat.completions.create(
                model=AI_MODEL,
                messages=retry_messages,
                temperature=0.1,
                max_tokens=4096,
            )
            retry_raw = retry_resp.choices[0].message.content.strip()
            parsed = _parse_json_payload(retry_raw)
        if not parsed:
            raise json.JSONDecodeError("invalid json payload", raw, 0)
        return _normalize_content_payload(
            parsed,
            theme=theme,
            phil=phil,
            activities=activities,
            child_initiative=child_initiative,
            child_desc=child_desc,
        )

    except json.JSONDecodeError as e:
        if not ALLOW_MOCK_CONTENT:
            logger.error("AI JSON 解析失败且禁止 Mock 回退：%s", e)
            raise HTTPException(
                status_code=502,
                detail="AI 返回内容格式异常（JSON 解析失败），请重试或检查模型配置。",
            )
        logger.warning("AI JSON 解析失败，ALLOW_MOCK_CONTENT=1，回退规范化 Mock：%s", e)
        return _normalize_content_payload(
            _mock_content(theme, phil, activities),
            theme=theme,
            phil=phil,
            activities=activities,
            child_initiative=child_initiative,
            child_desc=child_desc,
        )
    except Exception as e:
        _raise_if_invalid_dashscope_key(e)
        raise HTTPException(
            status_code=502,
            detail=f"调用阿里云百炼 API 失败：{e}",
        )


def _mock_content(theme: str, phil: str, activities: list[str]) -> dict:
    """开发调试用 Mock 数据（无需 API Key）"""
    return {
        "weekly_targets": {
            "teaching": f"围绕「{theme}」在五大领域开展均衡经验建构，体现{phil}特色。",
            "life": "提升卫生保健、安全与生活自理能力，形成稳定日常习惯。",
            "family": "明确家园配合重点，促进家庭端对主题学习的延伸支持。",
            "environment": f"依据「{theme}」配置环创与材料支持，如角色区与主题操作材料。",
        },
        "goals": [
            f"【健康领域】围绕「{theme}」，发展幼儿大肌肉协调能力与身体控制能力",
            f"【语言领域】通过「{theme}」情境丰富词汇，培养表达与倾听能力",
            f"【社会领域】在「{theme}」活动中建立合作意识与规则感",
            f"【科学领域】以「{theme}」为载体，激发观察与探究兴趣",
            f"【艺术领域】感受「{theme}」之美，发展创意表达能力",
        ],
        "preparation": [
            f"「{theme}」主题相关图片卡片与实物材料",
            "《3-6岁儿童学习与发展指南》领域目标对照单",
            "幼儿观察记录表",
        ],
        "activities": {a: f"围绕「{theme}」开展{a}（{phil}理念指导）" for a in activities},
        "child_initiative_note": "",
        "nap_guidance": "睡前提醒幼儿如厕与饮水；指导仰卧或侧卧，衣物分类摆放；起床后协助整理穿脱衣物，培养自理能力。",
        "evaluation": (
            f"通过观察幼儿在「{theme}」活动中的参与度、语言表达及合作行为，"
            f"参照{phil}理念的核心经验指标进行发展性评价。"
        ),
    }


# ──────────────────────────────────────────────
# Word 铁律填充
# ──────────────────────────────────────────────
ACTIVITY_LABEL_MAP = {
    "morning":     "🌅 晨间运动",
    "outdoor":     "🌿 户外活动",
    "environment": "🎨 环创活动",
    "life":        "🍽 生活活动",
    "area":        "🧩 区域活动",
    "family":      "👨‍👩‍👧 家园活动",
    "departure":   "🌙 离园活动",
}

WEEKDAY_TAGS = (
    ("星期一", "mon"),
    ("星期二", "tue"),
    ("星期三", "wed"),
    ("星期四", "thu"),
    ("星期五", "fri"),
    ("周一", "mon"),
    ("周二", "tue"),
    ("周三", "wed"),
    ("周四", "thu"),
    ("周五", "fri"),
)

FIVE_DOMAINS = ("健康", "语言", "社会", "科学", "艺术")

CLASS_LEVEL_HINTS: dict[str, str] = {
    "小班": (
        "【小班（3-4岁）特征】活动以感官体验、重复操作、生活自理为主；"
        "目标用词：感受、尝试、愿意、喜欢、在教师帮助下；"
        "活动时长短（10-15分钟）、材料大而安全、规则简单直接；"
        "户外侧重大肌肉运动（走跑跳爬）、区域侧重娃娃家与建构区。"
    ),
    "中班": (
        "【中班（4-5岁）特征】活动重探究尝试、规则合作、语言表达；"
        "目标用词：能够、学会、初步理解、主动参与、与同伴合作；"
        "活动时长中等（15-20分钟）、材料多样化、开始引入小组任务；"
        "户外增加器械组合与规则游戏、区域增加科学区与美工区深度。"
    ),
    "大班": (
        "【大班（5-6岁）特征】活动重深度探究、自主计划、问题解决、社会性成长；"
        "目标用词：自主、比较、发现规律、合作完成、独立表达观点；"
        "活动时长较长（20-30分钟）、材料低结构化、鼓励幼儿自主设计玩法；"
        "户外增加竞技合作与冒险挑战、区域强调项目式学习与跨区联动。"
    ),
}


def _weekday_tag_from_header(text: str) -> str | None:
    t = str(text or "").strip()
    t_compact = "".join(t.split())  # 去内部空白，兼容"星 期 一"
    for alias, tag in WEEKDAY_TAGS:
        if alias in t or alias in t_compact:
            return tag
    return None


def _build_weekday_domain_plan(theme: str) -> dict[str, str]:
    """
    周维度五大领域均衡分配：
    - 每周一到周五各有一个主领域
    - 主题不同会有不同起始位，避免始终固定同一顺序
    """
    seed = sum(ord(ch) for ch in str(theme or ""))
    start = seed % len(FIVE_DOMAINS)
    ordered = [FIVE_DOMAINS[(start + i) % len(FIVE_DOMAINS)] for i in range(5)]
    return {
        "mon": ordered[0],
        "tue": ordered[1],
        "wed": ordered[2],
        "thu": ordered[3],
        "fri": ordered[4],
    }


def _aw_doc_from_bytes(data: bytes) -> aw.Document:
    _aw_require()
    stream = io.BytesIO(data)
    return aw.Document(stream)


def _aw_doc_to_bytes(doc: aw.Document) -> bytes:
    _aw_require()
    out = io.BytesIO()
    doc.save(out, aw.SaveFormat.DOCX)
    return out.getvalue()


def _aw_require() -> None:
    if not ENABLE_ASPOSE_WORDS:
        raise RuntimeError("Aspose.Words 已关闭（设置 ENABLE_ASPOSE_WORDS=1 可启用）")
    if aw is None:
        _aw_lazy_import()
    if aw is None:
        raise RuntimeError("未安装 aspose-words，无法导出 Word（请检查 requirements / 容器镜像）")


def _aw_cell_dedup_key(cell) -> int:
    """一次填充会话内去重；Aspose Python 绑定未必提供 get_hash_code。"""
    return id(cell)


def _aw_node_as_table(node) -> aw.tables.Table:
    _aw_require()
    try:
        return node.as_table()
    except Exception:
        return aw.tables.Table.cast(node)


def _aw_stamp_export_provenance(doc: aw.Document) -> None:
    """
    在文档内写入可核对信息（Word：文件 → 信息 → 属性 → 备注/高级属性 → 自定义）。
    说明：正式授权后 Aspose 默认无评估水印；此处为应用层标记，便于确认走了 Aspose 管线。
    """
    try:
        doc.built_in_document_properties.comments = (
            f"智伴幼师导出 · v{APP_VERSION} · 排版引擎 Aspose.Words"
        )
        cdp = doc.custom_document_properties
        for name, val in (("ZhibanExportEngine", "Aspose.Words"), ("ZhibanAppVersion", APP_VERSION)):
            try:
                cdp.add(name, val)
            except Exception:
                # 已存在或 API 差异时忽略
                pass
    except Exception:
        pass


def _export_http_headers(engine: str) -> dict[str, str]:
    """engine: aspose-words | python-docx（与响应体实际使用的填充引擎一致）。"""
    return {
        "X-Export-Engine": engine,
        "X-App-Version": APP_VERSION,
    }


def _aw_cell_text(cell: aw.tables.Cell) -> str:
    try:
        return cell.to_string(aw.SaveFormat.TEXT).strip()
    except Exception:
        return ""


def _aw_cell_has_color(cell: aw.tables.Cell) -> bool:
    """
    Aspose 版着色检测：非透明/非白色底纹视为模板固定区。
    """
    try:
        shading = cell.cell_format.shading
        bg = shading.background_pattern_color.to_argb()
        fg = shading.foreground_pattern_color.to_argb()
        # 0=透明/空；-1 通常接近 auto。仅把显式白色排除。
        if bg not in (0, -1, 0x00FFFFFF, 0xFFFFFFFF):
            return True
        if fg not in (0, -1, 0x00FFFFFF, 0xFFFFFFFF):
            return True
    except Exception:
        return False
    return False


def _aw_copy_run_font(src_run: aw.Run, dst_run: aw.Run) -> None:
    try:
        dst_run.font.name = src_run.font.name
        dst_run.font.size = src_run.font.size
        dst_run.font.bold = src_run.font.bold
        dst_run.font.italic = src_run.font.italic
        dst_run.font.color = src_run.font.color
    except Exception:
        pass


def _aw_write_cell_preserve_style(doc: aw.Document, cell: aw.tables.Cell, text: str) -> None:
    """
    在 Aspose 中向单元格写入文本（多段换行）。
    兼容部分合并单元格结构：清空现有段落后再写入，避免部分 API 抛错导致导出失败。
    """
    _aw_require()
    lines = (text or "").split("\n")
    if not lines:
        lines = [""]
    while cell.paragraphs.count > 0:
        cell.paragraphs[0].remove()
    for line in lines:
        p = aw.Paragraph(doc)
        r = aw.Run(doc, line)
        p.append_child(r)
        cell.append_child(p)


def _aw_get_day_col_map(table: aw.tables.Table) -> tuple[int, list[int]]:
    weekdays = ("星期一", "星期二", "星期三", "星期四", "星期五", "周一", "周二", "周三", "周四", "周五")
    for ri in range(table.rows.count):
        row = table.rows[ri]
        day_cols: list[int] = []
        for ci in range(row.cells.count):
            t = _aw_cell_text(row.cells[ci])
            if any(d in t for d in weekdays):
                day_cols.append(ci)
        if len(day_cols) >= 3:
            return ri, sorted(day_cols)
    return -1, []


def _get_cell_text(cell) -> str:
    """获取单元格纯文本（去除空白）"""
    return "".join(p.text for p in cell.paragraphs).strip()


def _copy_run_format(src_run, dst_run):
    """将 src_run 的字体/颜色格式复制到 dst_run（不改变文字内容）"""
    if src_run.font.name:
        dst_run.font.name = src_run.font.name
        # 同步 East Asia 字体（中文环境必需）
        rPr = dst_run._r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:eastAsia"), src_run.font.name)
    if src_run.font.size:
        dst_run.font.size = src_run.font.size
    if src_run.font.bold is not None:
        dst_run.font.bold = src_run.font.bold
    if src_run.font.color and src_run.font.color.type is not None:
        try:
            dst_run.font.color.rgb = src_run.font.color.rgb
        except Exception:
            pass


def _get_day_col_map(table) -> tuple[int, list[int]]:
    """
    扫描表格，找到含 ≥3 个星期X / 周X 的行（星期表头行）及其列索引。
    返回 (day_header_row_idx, sorted_weekday_col_indices)；
    未找到则返回 (-1, [])。
    """
    weekdays = (
        "星期一", "星期二", "星期三", "星期四", "星期五",
        "周一", "周二", "周三", "周四", "周五",
    )
    for ri, row in enumerate(table.rows):
        cells = row.cells
        day_cols = [
            ci for ci, c in enumerate(cells)
            if any(d in "".join(_get_cell_text(c).split()) for d in weekdays)
        ]
        if len(day_cols) >= 3:
            return ri, sorted(day_cols)
    return -1, []


def _write_cell_preserve_style(cell, text: str) -> None:
    """
    向已净空的单元格写入文本：
    - 多行内容（\n 分隔）分别写入独立段落，不塞进单个 run
    - 从第一个 run 复制字体样式（净空引擎保留了格式壳）
    - 尽量复用现有段落结构，额外行才 add_paragraph
    """
    if not cell.paragraphs:
        for line in text.split("\n"):
            cell.add_paragraph(line)
        return

    lines = text.split("\n")

    # 找第一个 run 作为样式来源（净空后 run.text 为空但有 rPr）
    template_run = None
    for para in cell.paragraphs:
        for run in para.runs:
            template_run = run
            break
        if template_run:
            break

    # 清空所有 run 文本
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""

    # 逐行写入：复用已有段落，超出部分 add_paragraph
    for i, line in enumerate(lines):
        if i < len(cell.paragraphs):
            para = cell.paragraphs[i]
            if para.runs:
                para.runs[0].text = line
                if template_run and para.runs[0] is not template_run:
                    _copy_run_format(template_run, para.runs[0])
            else:
                r = para.add_run(line)
                if template_run:
                    _copy_run_format(template_run, r)
        else:
            new_p = cell.add_paragraph()
            r = new_p.add_run(line)
            if template_run:
                _copy_run_format(template_run, r)


def _build_weekly_fill_data(
    theme: str,
    phil: str,
    ai_content: dict,
    activities: list[str],
    child_initiative: bool,
    child_desc: str,
    *,
    class_level: str = "",
    fill_unselected: bool = False,
) -> dict[str, str]:
    """周计划填充字段字典（Aspose / python-docx 共用）。
    class_level: 班级型号 "小班" / "中班" / "大班"，影响集中活动规则。
    """
    today = _today_str()
    goals_text = "\n".join(
        f"目标{i+1}：{g}" for i, g in enumerate(ai_content.get("goals", []))
    )
    prep_text = "\n".join(
        f"• {p}" for p in ai_content.get("preparation", [])
    )
    weekly_targets = ai_content.get("weekly_targets", {}) if isinstance(ai_content, dict) else {}
    weekly_teaching = str((weekly_targets or {}).get("teaching", "") or "").strip()
    weekly_life = str((weekly_targets or {}).get("life", "") or "").strip()
    weekly_family = str((weekly_targets or {}).get("family", "") or "").strip()
    weekly_environment = str((weekly_targets or {}).get("environment", "") or "").strip()
    fill_data: dict[str, str] = {
        "theme":           theme,
        "philosophy":      phil,
        "goals":           goals_text,
        "preparation":     prep_text,
        "evaluation":      ai_content.get("evaluation", ""),
        "week_overview":   "、".join([ACTIVITY_LABEL_MAP.get(a, a) for a in activities]) or "本周按班级节奏灵活安排",
        "daily_points":    "",
        "key_difficulty":  weekly_teaching or "重点：围绕主题形成连续经验；难点：兼顾个体差异与活动节奏。",
        "children_baseline": child_desc if child_desc else "基于前期观察记录，幼儿对本周主题已有初步兴趣与经验基础。",
        "resource_list":   prep_text,
        "observation_plan": ai_content.get("evaluation", ""),
        "individual_support": ai_content.get("child_initiative_note", "") or "关注不同发展水平幼儿，提供分层支持与差异化引导。",
        "safety_risk":     "重点关注户外活动与材料使用安全；提前进行规则提醒与风险巡视。",
        "next_week_plan":  "根据本周观察结果调整下周材料投放与活动难度，延续幼儿高兴趣点。",
        "nap_guidance":    ai_content.get("nap_guidance", ""),
        "class_info":      f"________班    日期：{today}    天气：☀️ 晴",
        "child_initiative": (
            f"✅ 本周有幼儿自主发起活动\n{child_desc}\n"
            f"💡 {phil}理念中，幼儿自主发起的活动是最宝贵的课程生长点，请及时记录跟进。"
            if child_initiative and child_desc
            else ("✅ 本周有幼儿自主发起活动" if child_initiative else "")
        ),
    }
    if weekly_life:
        fill_data["life"] = weekly_life
    if weekly_family:
        fill_data["family"] = weekly_family
    if weekly_environment:
        fill_data["environment"] = weekly_environment
    weekday_domain = _build_weekday_domain_plan(theme)
    day_zh = {"mon": "周一", "tue": "周二", "wed": "周三", "thu": "周四", "fri": "周五"}
    fill_data["daily_points"] = "\n".join(
        f"{day_zh[tag]}（{weekday_domain.get(tag, '综合')}）：围绕「{theme}」推进核心经验，体现{phil}特色。"
        for tag in ("mon", "tue", "wed", "thu", "fri")
    )
    ai_acts: dict = ai_content.get("activities", {})
    for act_id in ACTIVITY_LABEL_MAP:
        content = ai_acts.get(act_id, "")
        if act_id in activities and content:
            fill_data[act_id] = content
        elif fill_unselected:
            fill_data[act_id] = "（本周未启用该板块，可按班级实际勾选后再生成）"
        base_for_day = fill_data.get(act_id, "") or content
        if base_for_day:
            for tag in ("mon", "tue", "wed", "thu", "fri"):
                domain = weekday_domain.get(tag, "综合")
                fill_data[f"{act_id}__{tag}"] = f"{base_for_day}\n【{day_zh[tag]}·{domain}】结合班级现状分层引导。"
    guidance_items = ai_content.get("guidance", [])
    if guidance_items:
        fill_data["guidance"] = "\n".join(
            f"{i+1}. {g}" for i, g in enumerate(guidance_items)
        )
    # study / game：AI 按"周X xxx\n周X xxx"格式返回，拆成每天独立内容写入对应格子
    # 【班级规则】小班：每天只取第一个活动；中班/大班：保留所有活动
    _day_prefix_re = re.compile(
        r"^(?:周[一二三四五]|星期[一二三四五])[^\S\n]*[：:·\-\s]*"
    )
    tag_order = ("mon", "tue", "wed", "thu", "fri")
    is_xiaopan = "小班" in class_level
    for field_id, field_default in (
        ("study", "本周集中活动（按班级课程表安排）"),
        ("game",  fill_data.get("area", "")),
    ):
        if field_id in fill_data:
            continue
        raw_text = ai_acts.get(field_id, "").strip()
        if not raw_text:
            fill_data[field_id] = field_default
            continue
        # 先按行拆，尝试匹配"周X"前缀，若能拆出 ≥2 行就写入 per-day 格子
        lines = [l.strip() for l in raw_text.splitlines() if l.strip()]
        per_day = [_day_prefix_re.sub("", l).strip() for l in lines]
        if len(per_day) >= 2:
            fill_data[field_id] = per_day[0]   # 基础 fallback 取第一天
            for i, tag in enumerate(tag_order):
                day_content = per_day[i] if i < len(per_day) else per_day[-1]
                # 【小班规则】：集中活动每天只有1个，取"；"前的部分
                if is_xiaopan and field_id == "study" and "；" in day_content:
                    day_content = day_content.split("；")[0].strip()
                fill_data[f"{field_id}__{tag}"] = day_content
        else:
            fill_data[field_id] = raw_text
    return fill_data


def _fill_word_template_docx_bytes(cleaned_bytes: bytes, fill_data: dict[str, str]) -> bytes:
    """python-docx 回填（Aspose 不可用或抛错时的可靠回退）。"""
    doc = Document(io.BytesIO(cleaned_bytes))
    filled_tc_elems: set = set()
    for table in doc.tables:
        rows = table.rows
        day_header_ri, day_cols = _get_day_col_map(table)
        if day_header_ri >= 0 and len(day_cols) >= 3:
            day_header_tc_elems: set = set()
            hdr_row = rows[day_header_ri]
            for day_ci in day_cols:
                if day_ci < len(hdr_row.cells):
                    day_header_tc_elems.add(hdr_row.cells[day_ci]._tc)
            header_cache: dict[int, str] = {}
            for day_ci in day_cols:
                if day_ci < len(hdr_row.cells):
                    header_cache[day_ci] = _get_cell_text(hdr_row.cells[day_ci])
            for row_idx in range(day_header_ri):
                cells = rows[row_idx].cells
                for col_idx, cell in enumerate(cells):
                    field = match_field(_get_cell_text(cell))
                    content = fill_data.get(field) if field else None
                    if not content:
                        continue
                    for ci in range(col_idx + 1, len(cells)):
                        cand = cells[ci]
                        if match_field(_get_cell_text(cand)) or _is_colored_cell(cand):
                            continue
                        tc_elem = cand._tc
                        if tc_elem in day_header_tc_elems:
                            continue
                        if tc_elem in filled_tc_elems:
                            break
                        filled_tc_elems.add(tc_elem)
                        _write_cell_preserve_style(cand, content)
                        break
            hdr_row = rows[day_header_ri]
            for day_ci, txt in header_cache.items():
                if day_ci < len(hdr_row.cells):
                    _write_cell_preserve_style(hdr_row.cells[day_ci], txt)
            current_field: str | None = None
            for row_idx in range(day_header_ri + 1, len(rows)):
                cells = rows[row_idx].cells
                if not cells:
                    continue
                col0_field = match_field(_get_cell_text(cells[0])) if _get_cell_text(cells[0]) else None
                if col0_field:
                    current_field = col0_field
                row_field = current_field
                if len(cells) > 1:
                    col1_text = _get_cell_text(cells[1])
                    col1_field = match_field(col1_text) if col1_text else None
                    if col1_field and col1_field in _ACTIVITY_FIELDS:
                        row_field = col1_field
                if not row_field:
                    continue
                base_content = fill_data.get(row_field)
                # Detect if all day columns point to the same merged cell;
                # if so, use base content (no per-day suffix — it's one shared cell).
                day_tc_ids = {
                    id(cells[ci]._tc)
                    for ci in day_cols if ci < len(cells)
                }
                all_days_merged = len(day_tc_ids) == 1
                for day_ci in day_cols:
                    if day_ci >= len(cells):
                        continue
                    target = cells[day_ci]
                    if _is_colored_cell(target):
                        continue
                    day_tag = _weekday_tag_from_header(header_cache.get(day_ci, ""))
                    # For merged single-cell rows use plain base content; otherwise
                    # prefer day-specific content (always populated regardless of selection).
                    if all_days_merged or not day_tag:
                        cell_content = base_content or fill_data.get(f"{row_field}__mon")
                    else:
                        cell_content = fill_data.get(f"{row_field}__{day_tag}") or base_content
                    if not cell_content:
                        continue
                    tc_elem = target._tc
                    if tc_elem in filled_tc_elems:
                        continue
                    filled_tc_elems.add(tc_elem)
                    _write_cell_preserve_style(target, cell_content)
        else:
            for row_idx, row in enumerate(rows):
                cells = row.cells
                for col_idx, cell in enumerate(cells):
                    field = match_field(_get_cell_text(cell))
                    content = fill_data.get(field) if field else None
                    if not content:
                        continue
                    target = None
                    if col_idx + 1 < len(cells):
                        cand = cells[col_idx + 1]
                        if not match_field(_get_cell_text(cand)) and not _is_colored_cell(cand):
                            target = cand
                    if target is None and row_idx + 1 < len(rows):
                        cand = rows[row_idx + 1].cells[col_idx]
                        if not match_field(_get_cell_text(cand)) and not _is_colored_cell(cand):
                            target = cand
                    if target is None:
                        continue
                    tc_elem = target._tc
                    if tc_elem in filled_tc_elems:
                        continue
                    filled_tc_elems.add(tc_elem)
                    _write_cell_preserve_style(target, content)
    try:
        doc.core_properties.comments = (
            f"智伴幼师导出 · v{APP_VERSION} · 排版引擎 python-docx（Aspose 回退）"
        )
    except Exception:
        pass
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


def _fill_word_template_aspose_bytes(cleaned_bytes: bytes, fill_data: dict[str, str]) -> bytes:
    """Aspose 回填；失败时由上层捕获并回退 docx。"""
    _aw_require()
    doc = _aw_doc_from_bytes(cleaned_bytes)
    filled_cells: set[int] = set()
    tables = doc.get_child_nodes(aw.NodeType.TABLE, True)

    for ti in range(tables.count):
        table = _aw_node_as_table(tables[ti])
        day_header_ri, day_cols = _aw_get_day_col_map(table)

        if day_header_ri >= 0 and len(day_cols) >= 3:
            header_cache: dict[int, str] = {}
            hdr_row = table.rows[day_header_ri]
            for day_ci in day_cols:
                if day_ci < hdr_row.cells.count:
                    header_cache[day_ci] = _aw_cell_text(hdr_row.cells[day_ci])

            for row_idx in range(day_header_ri):
                row = table.rows[row_idx]
                for col_idx in range(row.cells.count):
                    cell = row.cells[col_idx]
                    field = match_field(_aw_cell_text(cell))
                    content = fill_data.get(field) if field else None
                    if not content:
                        continue
                    for ci in range(col_idx + 1, row.cells.count):
                        cand = row.cells[ci]
                        if match_field(_aw_cell_text(cand)) or _aw_cell_has_color(cand):
                            continue
                        key = _aw_cell_dedup_key(cand)
                        if key in filled_cells:
                            break
                        filled_cells.add(key)
                        _aw_write_cell_preserve_style(doc, cand, content)
                        break

            hdr_row = table.rows[day_header_ri]
            for day_ci, txt in header_cache.items():
                if day_ci < hdr_row.cells.count:
                    _aw_write_cell_preserve_style(doc, hdr_row.cells[day_ci], txt)

            current_field: str | None = None
            for row_idx in range(day_header_ri + 1, table.rows.count):
                row = table.rows[row_idx]
                if row.cells.count == 0:
                    continue

                col0_field = match_field(_aw_cell_text(row.cells[0]))
                if col0_field:
                    current_field = col0_field
                row_field = current_field
                if row.cells.count > 1:
                    col1_field = match_field(_aw_cell_text(row.cells[1]))
                    if col1_field and col1_field in _ACTIVITY_FIELDS:
                        row_field = col1_field
                content = fill_data.get(row_field) if row_field else None
                if not content:
                    continue

                for day_ci in day_cols:
                    if day_ci >= row.cells.count:
                        continue
                    target = row.cells[day_ci]
                    if _aw_cell_has_color(target):
                        continue
                    day_tag = _weekday_tag_from_header(header_cache.get(day_ci, ""))
                    if day_tag and row_field:
                        content = fill_data.get(f"{row_field}__{day_tag}", content)
                    key = _aw_cell_dedup_key(target)
                    if key in filled_cells:
                        continue
                    filled_cells.add(key)
                    _aw_write_cell_preserve_style(doc, target, content)
        else:
            for row_idx in range(table.rows.count):
                row = table.rows[row_idx]
                for col_idx in range(row.cells.count):
                    cell = row.cells[col_idx]
                    field = match_field(_aw_cell_text(cell))
                    content = fill_data.get(field) if field else None
                    if not content:
                        continue

                    target = None
                    if col_idx + 1 < row.cells.count:
                        cand = row.cells[col_idx + 1]
                        if not match_field(_aw_cell_text(cand)) and not _aw_cell_has_color(cand):
                            target = cand
                    if target is None and row_idx + 1 < table.rows.count and col_idx < table.rows[row_idx + 1].cells.count:
                        cand = table.rows[row_idx + 1].cells[col_idx]
                        if not match_field(_aw_cell_text(cand)) and not _aw_cell_has_color(cand):
                            target = cand
                    if target is None:
                        continue
                    key = _aw_cell_dedup_key(target)
                    if key in filled_cells:
                        continue
                    filled_cells.add(key)
                    _aw_write_cell_preserve_style(doc, target, content)

    _aw_stamp_export_provenance(doc)
    return _aw_doc_to_bytes(doc)


def _build_content_disposition(filename: str) -> str:
    """
    构造兼容中英文文件名的下载头，避免 latin-1 编码错误导致 500。
    """
    filename = filename.strip() or "export.docx"
    utf8_name = quote(filename)
    ascii_fallback = re.sub(r"[^A-Za-z0-9._-]+", "_", filename)
    if not ascii_fallback:
        ascii_fallback = "export.docx"
    return f"attachment; filename=\"{ascii_fallback}\"; filename*=UTF-8''{utf8_name}"


def fill_word_template(
    template_bytes: bytes,
    theme: str,
    phil: str,
    ai_content: dict,
    activities: list[str],
    child_initiative: bool,
    child_desc: str,
    *,
    class_level: str = "",
    fill_unselected: bool = False,
) -> tuple[bytes, str]:
    """
    铁律填充：先净空模板骨架，再写入 AI 内容。
    优先 Aspose.Words；导入失败或运行异常时自动回退 python-docx，保证导出可用。
    返回 (docx 字节, 实际使用的引擎标识 aspose-words | python-docx)。
    """
    cleaned_bytes = clean_template_keep_style(template_bytes)
    fill_data = _build_weekly_fill_data(
        theme, phil, ai_content, activities, child_initiative, child_desc,
        class_level=class_level, fill_unselected=fill_unselected,
    )
    try:
        if ENABLE_ASPOSE_WORDS:
            return (
                _fill_word_template_aspose_bytes(cleaned_bytes, fill_data),
                "aspose-words",
            )
    except Exception as e:
        logger.warning("Aspose 周计划填充失败，回退 python-docx：%s", e, exc_info=True)
    return (
        _fill_word_template_docx_bytes(cleaned_bytes, fill_data),
        "python-docx",
    )


def docx_to_pdf_bytes(docx_bytes: bytes) -> bytes:
    """使用 Aspose.Words 将 DOCX 转换为 PDF。
    若 Aspose 未启用，抛出异常由上层捕获。"""
    _aw_require()
    doc = _aw_doc_from_bytes(docx_bytes)
    out = io.BytesIO()
    doc.save(out, _aw.SaveFormat.PDF)
    out.seek(0)
    return out.read()


def docx_to_images_bytes(docx_bytes: bytes, format: str = "png", dpi: int = 150) -> list[bytes]:
    """DOCX → PDF → 图片列表。每页一张图。format: 'png' 或 'jpg'。"""
    try:
        from pdf2image import convert_from_bytes
    except ImportError:
        raise RuntimeError("pdf2image 未安装，无法导出图片。请运行 pip install pdf2image")

    # 先转 PDF（用 Aspose，如果不可用则尝试系统 LibreOffice）
    try:
        pdf_bytes = docx_to_pdf_bytes(docx_bytes)
    except Exception:
        # 回退：尝试用 LibreOffice 命令行
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.write(docx_bytes)
            docx_path = f.name
        try:
            pdf_path = docx_path.replace(".docx", ".pdf")
            import subprocess
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf", docx_path],
                check=True, capture_output=True, timeout=30
            )
            with open(pdf_path, "rb") as f:
                pdf_bytes = f.read()
        finally:
            import os
            os.unlink(docx_path)
            if os.path.exists(pdf_path):
                os.unlink(pdf_path)

    # PDF → 图片列表
    images = convert_from_bytes(pdf_bytes, dpi=dpi, fmt=format)
    out_bytes = []
    for img in images:
        buf = io.BytesIO()
        img.save(buf, format=format.upper() if format != "jpg" else "JPEG")
        buf.seek(0)
        out_bytes.append(buf.read())
    return out_bytes


def _today_str() -> str:
    from datetime import date
    d = date.today()
    return f"{d.year}年{d.month}月{d.day}日"


def _build_standard_weekly_template_bytes() -> bytes:
    """生成标准周模板（19 模块超集，含可识别标签）。"""
    doc = Document()
    doc.add_heading("幼儿园周计划（标准模板 · 19模块）", level=1)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    fields = [
        ("基础信息", "园所：________    班级：________    周次：________    日期：________"),
        ("本周主题", ""),
        ("教育理念", ""),
        ("周总目标（五大领域）", ""),
        ("本周重点与难点", ""),
        ("幼儿已有经验", ""),
        ("本周活动总览表", ""),
        ("每日活动要点", ""),
        ("区域活动设计", ""),
        ("户外与体能活动", ""),
        ("生活活动与保育", ""),
        ("环境创设", ""),
        ("家园共育", ""),
        ("个别化支持", ""),
        ("安全与风险提示", ""),
        ("资源与材料清单", ""),
        ("观察记录计划", ""),
        ("周反思", ""),
        ("下周衔接", ""),
        ("午睡指导", ""),
        ("教师签名", "主班：________    配班：________"),
    ]
    for left, right in fields:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = right

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


def _build_standard_daily_template_bytes() -> bytes:
    """生成标准日教案空白模板（含可识别标签）。"""
    doc = Document()
    doc.add_heading("幼儿园日教案（标准模板）", level=1)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    fields = [
        ("活动名称", ""),
        ("教育理念", ""),
        ("班级", "________班    日期：________"),
        ("活动目标", ""),
        ("活动准备", ""),
        ("活动导入", ""),
        ("活动过程", ""),
        ("活动延伸", ""),
        ("活动反思", ""),
        ("观察要点", ""),
    ]
    for left, right in fields:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = right

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


WEEKDAY_HEADERS = ("星期一", "星期二", "星期三", "星期四", "星期五", "周一", "周二", "周三", "周四", "周五")
ROW_LABEL_HINTS = (
    "活动主题", "教学主题", "活动目标", "活动准备", "生活活动", "学习活动",
    "游戏活动", "区域活动", "户外活动", "家园活动", "离园活动", "评价与反思", "反思",
)


def _is_colored_cell(cell) -> bool:
    """
    判断单元格是否带底色（带底色视为模板固定区，不可动）。
    Word 常见着色入口：w:tcPr/w:shd @w:fill
    """
    try:
        tc_pr = cell._tc.tcPr
        if tc_pr is None:
            return False
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            return False
        fill = (shd.get(qn("w:fill")) or "").strip().upper()
        if not fill or fill in {"AUTO", "FFFFFF"}:
            return False
        return True
    except Exception:
        return False


def clean_template_keep_style(template_bytes: bytes) -> bytes:
    """
    将模板中“教师填写内容”净空，保留模板结构与样式。
    委托 kindergarten_template_cleaner（lxml 直接处理 document.xml）：
    色块 / 星期表头行 / 首列 / 模块关键词 / {{Tag}} 等规则与 TEMPLATE_STANDARD v1.1 同步。
    净空失败时回退原始字节，避免个别复杂模板导致整单导出 500。
    """
    try:
        return clean_docx_bytes(template_bytes)
    except Exception as e:
        logger.warning("净空模板失败，使用原始 .docx：%s", e, exc_info=True)
        return template_bytes


_PLACEHOLDER_RE = re.compile(r"\{\{\s*([^}]+?)\s*\}\}")


def analyze_template_docx(template_bytes: bytes) -> dict:
    """
    对上传模板做结构级自检，输出与 TEMPLATE_STANDARD.md v1.1 对齐的报告（启发式，非绝对准确率）。
    """
    doc = Document(io.BytesIO(template_bytes))
    day_header_rows = find_all_day_header_rows(template_bytes)
    table_stats: list[dict] = []
    all_text_joined: list[str] = []
    colored_cells = 0
    placeholder_tags: list[str] = []
    weekday_row_hits = 0
    field_hits: dict[str, int] = {}
    daily_hits: dict[str, int] = {}
    nested_label_rows = 0
    outdoor_nested_hint = False

    for ti, table in enumerate(doc.tables):
        rows = table.rows
        max_cols = 0
        for row in rows:
            max_cols = max(max_cols, len(row.cells))
        dhr = day_header_rows[ti] if ti < len(day_header_rows) else -1
        table_stats.append(
            {
                "table_index": ti + 1,
                "rows": len(rows),
                "max_cols": max_cols,
                "day_header_row": dhr,
            }
        )

        for row_idx, row in enumerate(rows):
            cells = row.cells
            first_text = _get_cell_text(cells[0]) if cells else ""
            if (
                first_text
                and any(h in first_text for h in ROW_LABEL_HINTS)
                and len(cells) >= 5
            ):
                nested_label_rows += 1

            for cell in cells:
                t = _get_cell_text(cell)
                if t:
                    all_text_joined.append(t)
                if _is_colored_cell(cell):
                    colored_cells += 1
                for m in _PLACEHOLDER_RE.finditer(t):
                    placeholder_tags.append(m.group(1).strip())
                for day in WEEKDAY_HEADERS:
                    if day in t:
                        weekday_row_hits += 1
                        break
                mf = match_field(t)
                if mf:
                    field_hits[mf] = field_hits.get(mf, 0) + 1
                md = match_daily_field(t)
                if md:
                    daily_hits[md] = daily_hits.get(md, 0) + 1

    blob = "".join(all_text_joined)
    if "户外" in blob and ("自主" in blob or "集体" in blob):
        outdoor_nested_hint = True

    # 文档类型启发式
    uniq_ph = list(dict.fromkeys(placeholder_tags))
    if uniq_ph:
        doc_type = "placeholder_template"
    elif len([k for k in daily_hits if k in ("introduction", "process", "extension", "reflection")]) >= 2:
        doc_type = "daily_plan"
    elif (
        ("星期一" in blob and "星期五" in blob)
        or ("周一" in blob and "周五" in blob)
        or weekday_row_hits >= 8
    ):
        doc_type = "weekly_grid"
    elif field_hits:
        doc_type = "general_plan"
    else:
        doc_type = "unknown"

    mappings: list[dict] = []
    if outdoor_nested_hint:
        subs = []
        if "自主" in blob:
            subs.append("自主")
        if "集体" in blob:
            subs.append("集体")
        mappings.append(
            {
                "tag": "plan.outdoor",
                "status": "partial_match",
                "reason": f"nested sub-rows hints: {subs}" if subs else "outdoor section present",
            }
        )

    score = 0.5
    if doc_type == "weekly_grid":
        score += 0.18
    elif doc_type == "daily_plan":
        score += 0.2
    elif doc_type == "placeholder_template":
        score += 0.22
    elif doc_type == "general_plan":
        score += 0.12
    if field_hits or daily_hits:
        score += 0.1
    if nested_label_rows:
        score -= 0.04
    score = round(min(0.95, max(0.4, score)), 2)

    return {
        "status": "success",
        "standard_ref": "TEMPLATE_STANDARD.md v1.2.0",
        "confidence_score": score,
        "doc_type_guess": doc_type,
        "document": {
            "tables": len(doc.tables),
            "day_header_rows": day_header_rows,
            "table_stats": table_stats,
            "colored_cells_detected": colored_cells,
            "placeholder_count": len(uniq_ph),
            "placeholders_sample": uniq_ph[:24],
            "weekday_cell_scan_hits": weekday_row_hits,
            "nested_wide_rows_hint": nested_label_rows,
        },
        "keyword_hits": {
            "fill_fields": field_hits,
            "daily_fields": daily_hits,
        },
        "mappings": mappings,
        "compliance": {
            "checklist": [
                "星期表头行（动态）与首列在净空引擎中默认保护；非星期表头的第一行按普通行处理",
                "下载「原文件名_净空模板.docx」与原文件并排对比表格是否变形",
                "若含 {{占位符}}，回填优先绑定占位符槽位；Word 可能拆分 run，检测需拼合文本",
            ],
        },
    }


# ──────────────────────────────────────────────
# 路由
# ──────────────────────────────────────────────
_FRONTEND = os.path.join(os.path.dirname(__file__), "index.html")


@app.get("/", tags=["前端"])
async def serve_frontend():
    """返回前端页面（同域部署时前后端共享一个 Cloud Run 服务）"""
    if os.path.exists(_FRONTEND):
        return FileResponse(_FRONTEND, media_type="text/html")
    return {
        "service": "智伴幼师 API",
        "status": "running",
        "model": AI_MODEL,
        "api_key_configured": bool(DASHSCOPE_API_KEY),
    }


@app.get("/health", tags=["健康检查"])
async def health():
    """Cloud Run / 负载均衡健康探针专用"""
    knowledge_base = _knowledge_base_status()
    return {
        "status": "ok",
        "app_version": APP_VERSION,
        "model": AI_MODEL,
        "api_key_configured": bool(DASHSCOPE_API_KEY),
        "api_key_note": "api_key_configured 仅表示已设置 DASHSCOPE_API_KEY，未校验密钥是否有效",
        "aspose_available": aw is not None,
        "export_note": (
            "/generate 与 /generate-daily 优先 Aspose，失败则回退 python-docx；"
            "响应头 X-Export-Engine 为本次实际使用的引擎"
        ),
        "knowledge_base": knowledge_base,
    }


@app.get("/knowledge-base/status", tags=["调试"])
async def knowledge_base_status():
    return {
        "status": "ok",
        "knowledge_base": _knowledge_base_status(),
    }


@app.post("/bug-report", tags=["调试"])
async def submit_bug_report(payload: dict = Body(...)):
    """
    前端导出失败上报入口：将上下文写入服务日志并返回 report_id。
    """
    from datetime import datetime
    import uuid

    report_id = f"BR-{datetime.utcnow().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}"
    safe = {
        "report_id": report_id,
        "created_at_utc": datetime.utcnow().isoformat() + "Z",
        "api_base": str(payload.get("api_base", ""))[:300],
        "page_url": str(payload.get("page_url", ""))[:800],
        "endpoint": str(payload.get("endpoint", ""))[:120],
        "error_message": str(payload.get("error_message", ""))[:2000],
        "theme": str(payload.get("theme", ""))[:200],
        "phil": str(payload.get("phil", ""))[:120],
        "file_name": str(payload.get("file_name", ""))[:240],
        "activities": payload.get("activities", []),
        "user_agent": str(payload.get("user_agent", ""))[:400],
        "app_version_front": str(payload.get("app_version_front", ""))[:40],
    }
    logger.error("EXPORT_BUG_REPORT %s", json.dumps(safe, ensure_ascii=False))
    return {"ok": True, "report_id": report_id}


@app.get("/public-stats", tags=["运营"])
async def public_stats():
    """首页公开统计：内测体验点击量、模块点击量、注册人数、建议留言数。"""
    stats = _load_app_stats()
    return {
        "status": "ok",
        "stats": {
            "home_visits": int(stats.get("home_visits", 0)),
            "module_clicks": int(stats.get("module_clicks", 0)),
            "register_count": int(stats.get("register_count", 0)),
            "feedback_count": int(stats.get("feedback_count", 0)),
        },
    }


@app.post("/track-event", tags=["运营"])
async def track_event(payload: dict = Body(...)):
    """
    轻量埋点：用于首页访问与模块点击计数。
    支持 event: home_visit | module_click | register
    """
    event = str(payload.get("event", "")).strip().lower()
    if event == "home_visit":
        stats = _inc_app_stat("home_visits", 1)
    elif event == "module_click":
        stats = _inc_app_stat("module_clicks", 1)
    elif event == "register":
        stats = _inc_app_stat("register_count", 1)
    else:
        raise HTTPException(status_code=400, detail="unsupported event")
    return {"ok": True, "stats": stats}


@app.post("/register-lite", tags=["运营"])
async def register_lite(payload: dict = Body(...)):
    """轻量内测注册（按联系方式去重计数）。"""
    identifier = str(payload.get("identifier", "")).strip().lower()
    role = str(payload.get("role", "")).strip().lower()
    if len(identifier) < 3:
        raise HTTPException(status_code=400, detail="请填写有效联系方式")
    if role not in {"teacher", "manager"}:
        raise HTTPException(status_code=400, detail="请先选择职业身份（老师/管理者）")

    existed = _load_registered_ids()
    is_new = identifier not in existed
    row = {
        "id": f"RG-{uuid4().hex[:10]}",
        "created_at_utc": _utc_iso(),
        "identifier": identifier[:160],
        "nickname": str(payload.get("nickname", "")).strip()[:80],
        "role": role,
        "kindergarten": str(payload.get("kindergarten", "")).strip()[:160],
        "page_url": str(payload.get("page_url", "")).strip()[:800],
    }
    _append_jsonl(_REGISTER_LOG_FILE, row)

    stats = _load_app_stats()
    if is_new:
        stats["register_count"] = int(stats.get("register_count", 0)) + 1
        _save_app_stats(stats)
    return {
        "ok": True,
        "registered": True,
        "is_new": bool(is_new),
        "register_count": int(_load_app_stats().get("register_count", 0)),
    }


@app.post("/feedback", tags=["运营"])
async def submit_feedback(payload: dict = Body(...)):
    """
    内测意见留言：前端提交建议后写入 jsonl 日志，并累计反馈数量。
    """
    from datetime import datetime
    import uuid

    message = str(payload.get("message", "")).strip()
    role = str(payload.get("role", "")).strip().lower()
    if role not in {"teacher", "manager"}:
        raise HTTPException(status_code=400, detail="请先选择职业身份（老师/管理者）")
    if len(message) < 2:
        raise HTTPException(status_code=400, detail="建议内容至少 2 个字")

    fb_id = f"FB-{datetime.utcnow().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:6]}"
    row = {
        "id": fb_id,
        "created_at_utc": datetime.utcnow().isoformat() + "Z",
        "nickname": str(payload.get("nickname", "")).strip()[:80],
        "contact": str(payload.get("contact", "")).strip()[:160],
        "role": role,
        "kindergarten": str(payload.get("kindergarten", "")).strip()[:160],
        "message": message[:3000],
        "page_url": str(payload.get("page_url", "")).strip()[:800],
        "module": str(payload.get("module", "")).strip()[:40],
        "user_agent": str(payload.get("user_agent", "")).strip()[:400],
    }
    try:
        with _FEEDBACK_LOG_FILE.open("a", encoding="utf-8") as f:
            f.write(json.dumps(row, ensure_ascii=False) + "\n")
    except Exception as e:
        logger.warning("写入反馈日志失败：%s", e)

    stats = _inc_app_stat("feedback_count", 1)
    logger.info("USER_FEEDBACK %s", json.dumps(row, ensure_ascii=False))
    return {
        "ok": True,
        "feedback_id": fb_id,
        "feedback_count": int(stats.get("feedback_count", 0)),
    }


register_redeem_routes(
    app,
    load_redeem_codes=_load_redeem_codes,
    redeem_code_core=_redeem_code_core,
    utc_iso=_utc_iso,
    dispatch_webhook=_dispatch_webhook,
    partner_redeem_tokens=PARTNER_REDEEM_TOKENS,
    partner_redeem_source=PARTNER_REDEEM_SOURCE,
    partner_webhook_urls=PARTNER_WEBHOOK_URLS,
    generate_unique_code=_generate_unique_code,
    save_redeem_codes=_save_redeem_codes,
    logger=logger,
)


@app.post("/user/wxlogin", tags=["用户"])
async def user_wxlogin(
    code: str = Form(..., description="微信 wx.login() 返回的 code"),
):
    """
    小程序登录接口：
    1. 用 code 换微信 openid（需配置 WECHAT_APPID + WECHAT_SECRET）
    2. 自动创建用户档案（含 agent_profile）
    3. 返回 user_token 用于后续所有请求认证
    """
    import httpx

    appid  = os.getenv("WECHAT_APPID", "")
    secret = os.getenv("WECHAT_SECRET", "")

    openid = ""
    if appid and secret:
        try:
            url = (
                f"https://api.weixin.qq.com/sns/jscode2session"
                f"?appid={appid}&secret={secret}&js_code={code}&grant_type=authorization_code"
            )
            async with httpx.AsyncClient(timeout=10) as client:
                resp = await client.get(url)
                wx_data = resp.json()
            openid = str(wx_data.get("openid", "")).strip()
            if not openid:
                raise HTTPException(status_code=400, detail=f"微信登录失败：{wx_data.get('errmsg', '未知')}")
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"微信接口调用失败：{e}")
    else:
        # 未配置微信密钥时，用 code 本身作为 openid（本地开发用）
        openid = f"dev_{code[:16]}"

    user = _get_or_create_user(openid)
    token = _generate_user_token(openid)

    # 写 token 到 Firestore + 本地
    accounts = _load_user_accounts()
    accounts[openid]["last_token"] = token
    accounts[openid]["last_login"] = _utc_iso()
    _save_user_accounts(accounts)

    return {
        "ok": True,
        "openid": openid,
        "user_token": token,
        "agent_profile": user.get("agent_profile", {}),
        "user_id": user.get("user_id") or "",
        "is_new": user.get("created_at") == accounts[openid].get("last_login"),
    }


@app.post("/user/agent", tags=["用户"])
async def update_agent_profile(
    user_token: str = Form(..., description="登录 token"),
    name: str = Form("小助手", description="智能体名字"),
    personality: str = Form("热心、耐心", description="性格特征"),
    tone: str = Form("亲切温暖", description="说话音调"),
    style: str = Form("鼓励式教学", description="教学风格"),
):
    """自定义用户的 AI 智能体性格，影响所有生成内容的风格。"""
    openid = _verify_user_token(user_token)
    accounts = _load_user_accounts()
    accounts[openid]["agent_profile"] = {
        "name": name,
        "personality": personality,
        "tone": tone,
        "style": style,
    }
    _save_user_accounts(accounts)
    return {"ok": True, "agent_profile": accounts[openid]["agent_profile"]}


@app.post("/user/register", tags=["用户"])
async def user_register(payload: dict = Body(...)):
    """
    极简内测注册：手机号即账号，重复注册视为登录，返回当前权益状态。
    """
    phone = str(payload.get("phone", "") or payload.get("user_id", "")).strip()
    if len(phone) < 5:
        raise HTTPException(status_code=400, detail="请填写有效手机号或账号（至少5位）")

    user_id = phone.lower()
    accounts = _load_user_accounts()
    is_new = user_id not in accounts
    now_iso = _utc_iso()

    if is_new:
        accounts[user_id] = {
            "user_id": user_id,
            "phone": phone,
            "created_at_utc": now_iso,
            "updated_at_utc": now_iso,
        }
    else:
        accounts[user_id]["updated_at_utc"] = now_iso
    _save_user_accounts(accounts)

    user_services = _load_user_services()
    service_entry = user_services.get(user_id, {})
    membership_until = service_entry.get("membership_until")
    is_active_member = False
    if membership_until:
        try:
            until_dt = datetime.fromisoformat(str(membership_until))
            is_active_member = datetime.now(timezone.utc) < until_dt
        except Exception:
            pass

    return {
        "ok": True,
        "is_new": is_new,
        "user_id": user_id,
        "created_at_utc": accounts[user_id]["created_at_utc"],
        "service": {
            "membership_until": membership_until,
            "is_active_member": is_active_member,
            "balance": int(service_entry.get("balance", 0) or 0),
            "quota": int(service_entry.get("quota", 0) or 0),
        },
    }


@app.get("/feedback-feed", tags=["运营"])
async def feedback_feed(limit: int = 30):
    """
    留言墙数据：默认匿名显示昵称，前端可点击“显示昵称”再展开。
    """
    safe_limit = max(1, min(int(limit or 30), 100))
    rows: list[dict] = []
    if _FEEDBACK_LOG_FILE.exists():
        try:
            lines = _FEEDBACK_LOG_FILE.read_text(encoding="utf-8").splitlines()
            for line in reversed(lines[-300:]):
                if len(rows) >= safe_limit:
                    break
                if not line.strip():
                    continue
                try:
                    raw = json.loads(line)
                except Exception:
                    continue
                nickname = str(raw.get("nickname", "")).strip()
                rows.append({
                    "id": str(raw.get("id", "")),
                    "created_at_utc": str(raw.get("created_at_utc", "")),
                    "role": str(raw.get("role", "")).strip().lower(),
                    "kindergarten": str(raw.get("kindergarten", "")).strip(),
                    "message": str(raw.get("message", "")).strip()[:500],
                    "nickname": nickname[:80],
                    "nickname_masked": "匿名用户",
                    "has_nickname": bool(nickname),
                })
        except Exception as e:
            logger.warning("读取反馈论坛失败：%s", e)
    return {"status": "ok", "items": rows}


@app.get("/template-standard", tags=["模板中心"])
async def template_standard():
    """返回程序内置模板识别标准（与 TEMPLATE_STANDARD 配置同步）。"""
    return {
        "status": "ok",
        "standard": TEMPLATE_STANDARD_V116,
    }


@app.get("/standard-templates", tags=["模板中心"])
async def get_standard_templates():
    """返回标准模板下载信息与累计下载次数。"""
    stats = _load_template_stats()
    return {
        "status": "ok",
        "templates": [
            {
                "id": "weekly",
                "name": "标准周/活动计划模板",
                "filename": "标准周活动计划模板.docx",
                "download_count": stats.get("weekly", 0),
            },
            {
                "id": "daily",
                "name": "标准日教案模板",
                "filename": "标准日教案模板.docx",
                "download_count": stats.get("daily", 0),
            },
            {
                "id": "cleaned",
                "name": "本次模板净空版",
                "filename": "本次模板净空版.docx",
                "download_count": stats.get("cleaned", 0),
            },
        ],
    }


@app.get("/standard-template/{template_id}/download", tags=["模板中心"])
async def download_standard_template(template_id: str):
    """下载标准模板，并记录下载次数。"""
    template_id = (template_id or "").strip().lower()
    if template_id == "weekly":
        data = _build_standard_weekly_template_bytes()
        filename = "标准周活动计划模板.docx"
    elif template_id == "daily":
        data = _build_standard_daily_template_bytes()
        filename = "标准日教案模板.docx"
    else:
        raise HTTPException(status_code=404, detail="未找到该标准模板")

    _inc_template_download(template_id)
    return StreamingResponse(
        io.BytesIO(data),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": _build_content_disposition(filename)},
    )


@app.post("/template/clean-download", tags=["模板中心"])
async def clean_download_template(
    template: UploadFile = File(..., description="老师上传的原始模板 .docx"),
):
    """
    基于老师上传模板生成“净空版标准模板”：
    - 仅删除教师填写/测试内容
    - 保留所有样式、间距、字体、表格结构
    """
    if not template.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="仅支持 .docx 格式")
    source = await template.read()
    if not source:
        raise HTTPException(status_code=400, detail="上传的文件为空")

    try:
        cleaned = clean_template_keep_style(source)
    except PackageNotFoundError:
        raise HTTPException(status_code=400, detail="模板解析失败，请确认 .docx 文件有效")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"模板净空失败：{e}")

    _inc_template_download("cleaned")
    original = os.path.basename(template.filename or "template.docx")
    original = re.sub(r'[\\/*?:"<>|]+', "_", original)
    if original.lower().endswith(".docx"):
        original = original[:-5]
    filename = f"{original}_净空模板.docx"

    return StreamingResponse(
        io.BytesIO(cleaned),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": _build_content_disposition(filename)},
    )


@app.post("/template/analyze", tags=["模板中心"])
async def template_analyze(
    template: UploadFile = File(..., description="待自检的老师模板 .docx"),
):
    """
    上传模板，返回与 TEMPLATE_STANDARD.md v1.1 对齐的自检 JSON（启发式）。
    用于合规验证：类型猜测、关键词命中、嵌套提示、置信度说明。
    """
    if not template.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="仅支持 .docx 格式")
    raw = await template.read()
    if not raw:
        raise HTTPException(status_code=400, detail="上传的文件为空")
    try:
        return analyze_template_docx(raw)
    except PackageNotFoundError:
        raise HTTPException(status_code=400, detail="模板解析失败，请确认 .docx 文件有效")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"模板分析失败：{e}")


@app.post("/generate", tags=["核心接口"])
async def generate(
    theme: str = Form(..., description="教学主题"),
    phil: str = Form(..., description="教育理念"),
    activities: str = Form("[]", description="活动重点列表（JSON 数组字符串）"),
    child_initiative: bool = Form(False, description="是否有幼儿自主发起活动"),
    child_desc: str = Form("", description="幼儿自主活动描述"),
    class_level: str = Form("", description="班级类型：小班 / 中班 / 大班"),
    client: str = Form("web", description="客户端标识，mini 时返回 JSON"),
    export_format: str = Form("docx", description="导出格式：docx / pdf / png"),
    template: UploadFile = File(..., description="Word 模板文件 (.doc/.docx)"),
):
    """
    核心生成接口

    - **theme**：教学主题，如「春天来了」
    - **phil**：教育理念，如「蒙氏教育（AMI/AMS）」
    - **activities**：活动重点，如 `["outdoor","area","morning"]`
    - **child_initiative**：是否有幼儿自主发起活动
    - **child_desc**：自主活动描述
    - **class_level**：班级类型（小班/中班/大班），影响活动难度与目标表述
    - **template**：用户上传的 Word 模板文件

    默认返回填充好的 `.docx` 文件流；
    当 `client=mini` 时，返回包含 `file_base64` 的 JSON，便于小程序直接写本地文件。
    """
    # ── 校验文件格式 ──
    if not template.filename.lower().endswith(".docx"):
        raise HTTPException(
            status_code=400,
            detail="仅支持 .docx 格式的 Word 文件",
        )

    # ── 解析活动列表 ──
    try:
        acts_list: list[str] = json.loads(activities)
    except (json.JSONDecodeError, TypeError):
        acts_list = []
    if not isinstance(acts_list, list):
        acts_list = []

    # ── 读取模板二进制 ──
    template_bytes = await template.read()
    if len(template_bytes) == 0:
        raise HTTPException(status_code=400, detail="上传的文件为空")
    _append_jsonl(_WEEKLY_DRAFT_LOG_FILE, {
        "ts": _utc_iso(),
        "event": "template_submit",
        "flow": "custom_template_generate",
        "template_filename": os.path.basename(template.filename or "template.docx"),
        "theme": theme,
        "phil": phil,
        "activities": acts_list,
        "child_initiative": bool(child_initiative),
    })

    # ── 调用 AI 生成内容 ──
    template_outline = _extract_template_outline(template_bytes)
    ai_content = generate_content(
        theme=theme,
        phil=phil,
        activities=acts_list,
        child_initiative=child_initiative,
        child_desc=child_desc,
        template_outline=template_outline,
        class_level=class_level.strip(),
    )

    # ── 铁律填充 ──
    try:
        filled_bytes, export_engine = fill_word_template(
            template_bytes=template_bytes,
            theme=theme,
            phil=phil,
            ai_content=ai_content,
            activities=acts_list,
            child_initiative=child_initiative,
            child_desc=child_desc,
            class_level=class_level.strip(),
        )
    except PackageNotFoundError:
        raise HTTPException(status_code=400, detail="模板解析失败，请确认上传的是有效 .docx 文件")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"模板填充失败：{e}")

    # ── 导出文件名 + 格式转换 ──
    original_name = os.path.basename(template.filename or "template.docx")
    original_name = re.sub(r'[\\/*?:"<>|]+', "_", original_name)
    if original_name.lower().endswith(".docx"):
        original_name = original_name[:-5]

    export_fmt = export_format.strip().lower()

    # 小程序返回 JSON
    if client.strip().lower() == "mini":
        return _build_mini_doc_payload(
            filled_bytes=filled_bytes,
            original_name=f"{original_name}.docx",
            export_engine=export_engine,
        )

    # 导出格式选择
    if export_fmt == "pdf":
        try:
            output_bytes = docx_to_pdf_bytes(filled_bytes)
            filename = f"{original_name}.pdf"
            media_type = "application/pdf"
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"PDF 生成失败：{e}")
    elif export_fmt == "png":
        try:
            images_bytes = docx_to_images_bytes(filled_bytes, format="png")
            if not images_bytes:
                raise ValueError("无法转换图片")
            # 多页则返回第一页，或在响应中打包所有页面
            # 简化版：返回第一页
            output_bytes = images_bytes[0]
            filename = f"{original_name}_page1.png"
            media_type = "image/png"
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"图片导出失败：{e}")
    else:  # docx（默认）
        output_bytes = filled_bytes
        filename = f"{original_name}.docx"
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    h = {
        "Content-Disposition": _build_content_disposition(filename),
        "X-AI-Model": AI_MODEL,
    }
    h.update(_export_http_headers(export_engine))
    return StreamingResponse(
        io.BytesIO(output_bytes),
        media_type=media_type,
        headers=h,
    )


# ──────────────────────────────────────────────
# 预览接口（仅返回 AI 内容 JSON，不需要模板）
# ──────────────────────────────────────────────
@app.post("/preview", tags=["调试"])
async def preview(
    theme: str = Form(...),
    phil: str = Form(...),
    activities: str = Form("[]"),
    child_initiative: bool = Form(False),
    child_desc: str = Form(""),
    class_level: str = Form("", description="班级类型：小班 / 中班 / 大班"),
):
    """返回 AI 生成的结构化内容 JSON，用于前端预览和调试（无需上传文件）"""
    acts_list: list[str] = json.loads(activities) if activities else []
    content = generate_content(theme, phil, acts_list, child_initiative, child_desc, class_level=class_level.strip())
    return {"status": "ok", "theme": theme, "philosophy": phil, "content": content}


@app.post("/standard-weekly-draft", tags=["标准周模板"])
async def standard_weekly_draft(
    theme: str = Form(...),
    phil: str = Form(...),
    activities: str = Form("[]"),
    child_initiative: bool = Form(False),
    child_desc: str = Form(""),
    class_level: str = Form("", description="班级类型：小班 / 中班 / 大班"),
):
    """生成标准周模板草稿（JSON，可在前端修改后再导出）。"""
    try:
        acts_list: list[str] = json.loads(activities) if activities else []
    except Exception:
        acts_list = []
    if not isinstance(acts_list, list):
        acts_list = []
    content = generate_content(theme, phil, acts_list, child_initiative, child_desc, class_level=class_level.strip())
    draft_id = f"wd_{uuid4().hex[:10]}"
    _append_jsonl(_WEEKLY_DRAFT_LOG_FILE, {
        "ts": _utc_iso(),
        "event": "standard_weekly_draft",
        "draft_id": draft_id,
        "theme": theme,
        "phil": phil,
        "activities": acts_list,
        "child_initiative": bool(child_initiative),
    })
    return {
        "status": "ok",
        "draft_id": draft_id,
        "schema_version": "weekly-standard-v1.0",
        "module_catalog": WEEKLY_STANDARD_MODULES,
        "draft": {
            "theme": theme,
            "phil": phil,
            "activities": acts_list,
            "child_initiative": bool(child_initiative),
            "child_desc": child_desc or "",
            "content": content,
        },
    }


@app.post("/standard-weekly-export", tags=["标准周模板"])
async def standard_weekly_export(
    draft_json: str = Form(..., description="前端可编辑草稿 JSON"),
):
    """按可编辑草稿导出标准周模板 Word。"""
    try:
        payload = json.loads(draft_json)
        if not isinstance(payload, dict):
            raise ValueError("draft_json 必须是对象")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"draft_json 解析失败：{e}")

    theme = str(payload.get("theme") or "").strip()
    phil = str(payload.get("phil") or "五大领域").strip()
    child_desc = str(payload.get("child_desc") or "").strip()
    child_initiative = bool(payload.get("child_initiative", False))
    acts_list = payload.get("activities", [])
    if not isinstance(acts_list, list):
        acts_list = []
    raw_content = payload.get("content", {})
    if not theme:
        raise HTTPException(status_code=400, detail="草稿缺少主题（theme）")

    content = _normalize_content_payload(
        raw_content if isinstance(raw_content, dict) else {},
        theme=theme,
        phil=phil,
        activities=acts_list,
        child_initiative=child_initiative,
        child_desc=child_desc,
    )
    template_bytes = _build_standard_weekly_template_bytes()
    try:
        filled_bytes, export_engine = fill_word_template(
            template_bytes=template_bytes,
            theme=theme,
            phil=phil,
            ai_content=content,
            activities=acts_list,
            child_initiative=child_initiative,
            child_desc=child_desc,
            fill_unselected=True,
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"标准模板导出失败：{e}")

    draft_id = str(payload.get("draft_id") or "")
    _append_jsonl(_WEEKLY_DRAFT_LOG_FILE, {
        "ts": _utc_iso(),
        "event": "standard_weekly_export",
        "draft_id": draft_id or f"wd_{uuid4().hex[:8]}",
        "theme": theme,
        "phil": phil,
        "activities": acts_list,
    })

    safe_theme = re.sub(r'[\\/*?:"<>|]+', "_", theme)
    filename = f"标准周模板_{safe_theme}.docx"
    h = {
        "Content-Disposition": _build_content_disposition(filename),
        "X-AI-Model": AI_MODEL,
    }
    h.update(_export_http_headers(export_engine))
    return StreamingResponse(
        io.BytesIO(filled_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=h,
    )


# ══════════════════════════════════════════════════════════════════════
#  周日联动模块
#  /generate-weekly  →  五天周计划 JSON
#  /generate-daily   →  单天四维日教案（导入/过程/延伸/反思）→ Word
# ══════════════════════════════════════════════════════════════════════

# ── 日教案单元格关键字映射（独立于活动计划表，针对日教案格式）──────
DAILY_CELL_KEYWORD_MAP: list[tuple[list[str], str]] = [
    (["活动名称", "课题名称", "活动标题", "课题"],   "title"),
    (["教育理念", "课程理念", "理念"],               "philosophy"),
    (["班级", "年龄段", "年龄", "班"],               "class_info"),
    (["活动目标", "教学目标", "目标"],               "goals"),
    (["活动准备", "材料准备", "准备"],               "preparation"),
    (["活动导入", "情境导入", "导入", "引入"],        "introduction"),
    (["活动过程", "基本过程", "主要过程", "过程"],    "process"),
    (["活动延伸", "延伸活动", "延伸", "拓展"],       "extension"),
    (["活动反思", "教师反思", "反思", "评价小结"],    "reflection"),
    (["观察要点", "观察记录", "观察重点"],            "observation"),
    (["日期", "时间"],                               "date_info"),
]


def match_daily_field(cell_text: str) -> Optional[str]:
    t = cell_text.strip()
    for keywords, field in DAILY_CELL_KEYWORD_MAP:
        for kw in keywords:
            if kw in t:
                return field
    return None


# ── 周计划 AI 生成 ──────────────────────────────────────────────────
def build_weekly_prompt(
    theme: str, phil: str, activities: list[str], class_level: str = ""
) -> str:
    """
    使用 Prompt 工程系统生成周计划 Prompt。

    这个函数调用 PromptTemplate 来确保生成的内容风格一致、质量稳定。
    可复现的：任何电脑、任何时间都产生相同质量的输出。
    """
    prompt_template = get_prompt_template()
    return prompt_template.build_user_prompt(
        theme=theme,
        class_level=class_level or "中班",  # 默认中班
        philosophy=phil,
        activities=activities or ["区域活动", "户外活动"],
    )


def generate_weekly_content(
    theme: str, phil: str, activities: list[str], class_level: str = ""
) -> dict:
    if not DASHSCOPE_API_KEY:
        return _mock_weekly(theme, phil)
    try:
        prompt_template = get_prompt_template()
        resp = client.chat.completions.create(
            model=AI_MODEL,
            messages=[
                {"role": "system", "content": prompt_template.build_system_prompt()},
                {
                    "role": "user",
                    "content": build_weekly_prompt(theme, phil, activities, class_level),
                },
            ],
            temperature=0.2,  # 降低温度以保持风格一致性
            max_tokens=4096,  # 足够容纳完整的周计划
        )
        raw = resp.choices[0].message.content.strip()
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)
        return json.loads(raw)
    except Exception as e:
        _raise_if_invalid_dashscope_key(e)
        raise HTTPException(status_code=502, detail=f"周计划生成失败：{e}")


def _mock_weekly(theme: str, phil: str) -> dict:
    tasks = [
        ("观察与记录", "感官探索", "区域活动"),
        ("创意表达", "艺术创作", "环创活动"),
        ("合作探究", "社交能力", "户外活动"),
        ("生活技能", "自理能力", "生活活动"),
        ("成果分享", "语言表达", "家园活动"),
    ]
    return {
        "week_theme": theme,
        "philosophy": phil,
        "days": [
            {
                "day": f"周{c}",
                "task": f"{theme}·{t}",
                "focus": f,
                "activity_type": at,
                "hint": f"关注幼儿在{f}方面的表现",
            }
            for (t, f, at), c in zip(tasks, ["一", "二", "三", "四", "五"])
        ],
    }


# ── 园部计划骨架（学期→月→周）────────────────────────────────────────
def _build_term_month_week_skeleton(
    term_theme: str,
    start_month: int,
    month_count: int,
) -> dict:
    """生成园部学期/月/周骨架，便于园所先统一规划，再下钻到班级执行。"""
    month_count = max(1, min(month_count, 6))
    months: list[dict] = []
    for i in range(month_count):
        month_no = ((start_month - 1 + i) % 12) + 1
        month_title = f"{month_no}月"
        weeks = [
            {
                "week_index": w + 1,
                "week_theme": f"{term_theme}·{month_title}第{w + 1}周",
                "focus": "园所统一目标待补充",
                "status": "planned",
            }
            for w in range(4)
        ]
        months.append(
            {
                "month": month_title,
                "month_goal": "围绕学期目标分解月度重点",
                "weeks": weeks,
            }
        )
    return {
        "term_theme": term_theme,
        "months": months,
    }


# ── 日教案 AI 生成 ──────────────────────────────────────────────────
def build_daily_prompt(
    week_theme: str, day: str, task: str, phil: str, phil_hint: str
) -> str:
    return textwrap.dedent(f"""
        请为幼儿园日教案生成专业内容，以 JSON 格式返回。

        【基本信息】
        - 周主题：{week_theme}
        - 今日活动（{day}）：{task}
        - 教育理念：{phil}

        【理念专业词汇要求】
        {phil_hint}

        【四维结构要求】
        请严格按照「导入 → 过程 → 延伸 → 反思」四个维度设计，每个维度字数控制在100-150字。

        【JSON 输出格式（只返回 JSON）】
        {{
          "title": "活动名称（即今日任务）",
          "goals": [
            "目标1（含维度前缀，如感知/语言/社会等）",
            "目标2",
            "目标3"
          ],
          "preparation": ["材料1", "材料2", "材料3"],
          "introduction": "导入环节（5-8分钟）：情境创设、激发兴趣、连接已有经验的具体步骤",
          "process": "活动过程（20-25分钟）：分步骤描述教师引导动作、幼儿操作内容、关键问题设计",
          "extension": "延伸活动（5-10分钟）：区域延伸、家园延伸或跨日连接建议",
          "reflection": "教师反思：今日活动的观察重点与课后反思问题（含{phil}理念专业术语）",
          "observation": "重点观察要点（3条，每条15字以内）"
        }}
    """).strip()


def generate_daily_content(
    week_theme: str, day: str, task: str, phil: str
) -> dict:
    if not DASHSCOPE_API_KEY:
        return _mock_daily(week_theme, day, task, phil)
    phil_hint = PHILOSOPHY_HINTS.get(phil, "")
    try:
        resp = client.chat.completions.create(
            model=AI_MODEL,
            messages=[
                {"role": "system", "content": build_system_prompt()},
                {"role": "user",   "content": build_daily_prompt(week_theme, day, task, phil, phil_hint)},
            ],
            temperature=0.7,
            max_tokens=1500,
        )
        raw = resp.choices[0].message.content.strip()
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)
        return json.loads(raw)
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=502, detail=f"日教案内容解析失败：{e}")
    except Exception as e:
        _raise_if_invalid_dashscope_key(e)
        raise HTTPException(status_code=502, detail=f"日教案生成失败：{e}")


def _mock_daily(week_theme: str, day: str, task: str, phil: str) -> dict:
    return {
        "title": task,
        "goals": [
            f"感知「{task}」的核心特征，积累相关直接经验",
            f"通过操作与探索，发展观察比较能力",
            f"在互动中提升语言表达与合作意识",
        ],
        "preparation": [f"「{task}」相关操作材料", "观察记录单", "展示板"],
        "introduction": (
            f"（{day}·导入）以情境导入：展示与「{task}」相关的真实物品或图片，"
            f"提问：'你在哪里见过这个？它让你想到了什么？'激活幼儿先备经验，"
            f"自然过渡到今日探究活动。（{phil}理念指导：注重儿童已有经验的联结）"
        ),
        "process": (
            f"（过程）①教师示范操作步骤，明确今日探究任务；"
            f"②幼儿自主操作，教师巡回观察并记录关键行为；"
            f"③小组分享：'你发现了什么？你是怎么做到的？'；"
            f"④集体总结，提炼「{task}」的核心经验。"
        ),
        "extension": (
            f"（延伸）区域延伸：在科学探究区/美工区延续「{task}」相关材料供幼儿自主探索；"
            f"家园延伸：请家长与孩子在家寻找与「{task}」相关的生活场景并拍照分享。"
        ),
        "reflection": (
            f"（{phil}·反思）观察今日幼儿在「{task}」活动中的参与度与深度思考迹象；"
            f"记录有价值的幼儿语言；思考：哪些材料激发了更持久的探究？"
            f"明日如何在此基础上推进？"
        ),
        "observation": "观察专注投入时长 | 记录语言表达关键词 | 关注合作协商行为",
    }


def _build_daily_fill_data(
    daily_content: dict,
    week_theme: str,
    day: str,
    phil: str,
) -> dict[str, str]:
    today = _today_str()
    goals_text = "\n".join(
        f"目标{i+1}：{g}" for i, g in enumerate(daily_content.get("goals", []))
    )
    prep_text = "\n".join(f"• {p}" for p in daily_content.get("preparation", []))
    obs_raw = daily_content.get("observation", "")
    obs_text = "\n".join(
        f"△ {o.strip()}"
        for o in (obs_raw.split("|") if "|" in obs_raw else [obs_raw])
        if o.strip()
    )
    return {
        "title":        daily_content.get("title", f"{week_theme}·{day}"),
        "philosophy":   phil,
        "class_info":   f"________班    日期：{today}",
        "goals":        goals_text,
        "preparation":  prep_text,
        "introduction": daily_content.get("introduction", ""),
        "process":      daily_content.get("process", ""),
        "extension":    daily_content.get("extension", ""),
        "reflection":   daily_content.get("reflection", ""),
        "observation":  obs_text,
        "date_info":    f"{today}（{day}）",
    }


def _build_daily_structured_docx_bytes(
    daily_content: dict,
    week_theme: str,
    day: str,
    phil: str,
    source_day: Optional[dict] = None,
) -> bytes:
    """
    生成“非表格”日计划文档：完整日结构 + 文末「生成」「调整」两个部分。
    """
    doc = Document()
    today = _today_str()
    title = daily_content.get("title", f"{week_theme}·{day}日计划")
    doc.add_heading(f"{title}（{day}）", level=1)

    intro = doc.add_paragraph()
    intro.add_run("周主题：").bold = True
    intro.add_run(f"{week_theme}    ")
    intro.add_run("日期：").bold = True
    intro.add_run(f"{today}    ")
    intro.add_run("教育理念：").bold = True
    intro.add_run(phil)

    goals = daily_content.get("goals", []) or []
    prep = daily_content.get("preparation", []) or []
    obs_raw = daily_content.get("observation", "") or ""
    obs_items = [x.strip() for x in (obs_raw.split("|") if "|" in obs_raw else [obs_raw]) if x.strip()]

    doc.add_heading("一、活动目标", level=2)
    if goals:
        for g in goals:
            doc.add_paragraph(str(g), style="List Bullet")
    else:
        doc.add_paragraph("（待补充）")

    doc.add_heading("二、活动准备", level=2)
    if prep:
        for p in prep:
            doc.add_paragraph(str(p), style="List Bullet")
    else:
        doc.add_paragraph("（待补充）")

    doc.add_heading("三、活动导入", level=2)
    doc.add_paragraph(daily_content.get("introduction", "（待补充）"))

    doc.add_heading("四、活动过程", level=2)
    doc.add_paragraph(daily_content.get("process", "（待补充）"))

    doc.add_heading("五、活动延伸", level=2)
    doc.add_paragraph(daily_content.get("extension", "（待补充）"))

    doc.add_heading("六、活动反思", level=2)
    doc.add_paragraph(daily_content.get("reflection", "（待补充）"))

    doc.add_heading("七、重点观察", level=2)
    if obs_items:
        for o in obs_items:
            doc.add_paragraph(o, style="List Bullet")
    else:
        doc.add_paragraph("（待补充）")

    doc.add_heading("八、生成", level=2)
    gen = doc.add_paragraph()
    gen.add_run("来源：").bold = True
    gen.add_run("周计划联动自动生成\n")
    gen.add_run("模型：").bold = True
    gen.add_run(f"{AI_MODEL}\n")
    gen.add_run("生成时间：").bold = True
    gen.add_run(f"{today}\n")
    if source_day:
        source_summary = (
            f"任务：{source_day.get('task', '')}；"
            f"关注点：{source_day.get('focus', '')}；"
            f"活动类型：{source_day.get('activity_type', '')}；"
            f"提示：{source_day.get('hint', '')}"
        )
        doc.add_paragraph(source_summary)

    doc.add_heading("九、调整", level=2)
    doc.add_paragraph("教师二次调整记录：")
    doc.add_paragraph("1. 今日微调点：________________________________________")
    doc.add_paragraph("2. 幼儿响应观察：______________________________________")
    doc.add_paragraph("3. 明日延伸计划：______________________________________")

    try:
        doc.core_properties.comments = (
            f"智伴幼师导出 · v{APP_VERSION} · 排版引擎 python-docx（结构化日计划）"
        )
    except Exception:
        pass
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


def _build_observation_prompt(
    theme: str,
    child_name: str,
    scene: str,
    note: str,
    phil: str,
    photo_names: list[str],
) -> str:
    photo_text = "、".join(photo_names) if photo_names else "（未上传照片）"
    child = child_name or "幼儿"
    return textwrap.dedent(
        f"""
        请为幼儿园教师生成一份「拍照观察记录」，只返回 JSON，不要其他文字。

        【输入信息】
        - 观察主题：{theme}
        - 观察对象：{child}
        - 观察场景：{scene}
        - 教育理念：{phil}
        - 照片文件名：{photo_text}
        - 教师补充：{note or "无"}

        【输出要求】
        - 语气专业、真实、可执行，避免空话
        - 贴近一线幼师记录语境
        - 每条建议可直接用于复盘和家园沟通

        【JSON 格式】
        {{
          "title": "观察记录标题",
          "summary": "观察概述（80-120字）",
          "records": ["关键观察1", "关键观察2", "关键观察3"],
          "analysis": "发展解读（80-120字）",
          "supports": ["支持策略1", "支持策略2", "支持策略3"],
          "home_cooperation": "家园共育建议（40-80字）",
          "next_plan": "下次跟进计划（40-80字）",
          "generated": "系统生成说明（简短）",
          "adjustment": "教师可调整建议（简短）"
        }}
        """
    ).strip()


def _mock_observation_content(
    theme: str,
    child_name: str,
    scene: str,
    note: str,
    photo_names: list[str],
) -> dict:
    child = child_name or "幼儿"
    topic = theme or "主题活动"
    note_text = f"教师补充：{note}" if note else "教师补充：无"
    return {
        "title": f"{topic}观察记录",
        "summary": (
            f"在{scene}中，{child}围绕「{topic}」表现出较高参与度，"
            "能够主动回应任务并与同伴互动。"
            "从现场照片与过程记录看，幼儿有持续投入与表达意愿。"
        ),
        "records": [
            f"{child}在活动中能主动操作材料，并保持阶段性专注。",
            "幼儿在同伴互动中出现协商与轮流行为，社交参与较积极。",
            "面对任务变化时，幼儿愿意尝试不同方法并表达自己的发现。",
        ],
        "analysis": (
            f"{child}在观察表达与合作参与方面已有可见进步。"
            "建议继续通过同类情境巩固经验，提升语言组织与问题解决深度。"
        ),
        "supports": [
            "提供可重复操作的低结构材料，支持幼儿再次探索。",
            "教师使用开放式提问，引导幼儿描述“如何发现”。",
            "设计双人协作小任务，强化沟通与角色分工体验。",
        ],
        "home_cooperation": "建议家长在家庭场景延续同主题观察活动，记录孩子的关键表达并反馈给教师。",
        "next_plan": "下次活动增加“分享与复述”环节，帮助幼儿梳理观察过程并形成表达闭环。",
        "generated": f"已接收 {len(photo_names)} 张照片；{note_text}",
        "adjustment": "教师可根据班级节奏对难度、材料与提问方式做个性化微调。",
    }


def generate_observation_content(
    theme: str,
    child_name: str,
    scene: str,
    note: str,
    phil: str,
    photo_names: list[str],
) -> dict:
    if not DASHSCOPE_API_KEY:
        return _mock_observation_content(theme, child_name, scene, note, photo_names)
    try:
        resp = client.chat.completions.create(
            model=AI_MODEL,
            messages=[
                {"role": "system", "content": build_system_prompt()},
                {
                    "role": "user",
                    "content": _build_observation_prompt(
                        theme=theme,
                        child_name=child_name,
                        scene=scene,
                        note=note,
                        phil=phil,
                        photo_names=photo_names,
                    ),
                },
            ],
            temperature=0.65,
            max_tokens=1200,
        )
        raw = resp.choices[0].message.content.strip()
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)
        return json.loads(raw)
    except Exception as e:
        _raise_if_invalid_dashscope_key(e)
        return _mock_observation_content(theme, child_name, scene, note, photo_names)


def _build_observation_docx_bytes(
    content: dict,
    theme: str,
    child_name: str,
    scene: str,
    phil: str,
    note: str,
    photo_names: list[str],
) -> bytes:
    doc = Document()
    today = _today_str()
    child = child_name or "幼儿"
    doc.add_heading(content.get("title", f"{theme}观察记录"), level=1)

    intro = doc.add_paragraph()
    intro.add_run("日期：").bold = True
    intro.add_run(f"{today}    ")
    intro.add_run("观察场景：").bold = True
    intro.add_run(f"{scene}    ")
    intro.add_run("观察对象：").bold = True
    intro.add_run(child)

    doc.add_paragraph(f"教育理念：{phil}")
    doc.add_paragraph(f"照片数量：{len(photo_names)}")
    if photo_names:
        doc.add_paragraph("照片文件：")
        for name in photo_names:
            doc.add_paragraph(name, style="List Bullet")

    if note:
        doc.add_heading("教师补充", level=2)
        doc.add_paragraph(note)

    doc.add_heading("一、观察概述", level=2)
    doc.add_paragraph(content.get("summary", "（待补充）"))

    doc.add_heading("二、关键观察记录", level=2)
    records = content.get("records", []) or []
    if records:
        for row in records:
            doc.add_paragraph(str(row), style="List Bullet")
    else:
        doc.add_paragraph("（待补充）")

    doc.add_heading("三、发展解读", level=2)
    doc.add_paragraph(content.get("analysis", "（待补充）"))

    doc.add_heading("四、支持策略", level=2)
    supports = content.get("supports", []) or []
    if supports:
        for row in supports:
            doc.add_paragraph(str(row), style="List Bullet")
    else:
        doc.add_paragraph("（待补充）")

    doc.add_heading("五、家园共育建议", level=2)
    doc.add_paragraph(content.get("home_cooperation", "（待补充）"))

    doc.add_heading("六、下次跟进计划", level=2)
    doc.add_paragraph(content.get("next_plan", "（待补充）"))

    doc.add_heading("七、生成", level=2)
    doc.add_paragraph(content.get("generated", "由系统生成初稿"))

    doc.add_heading("八、调整", level=2)
    doc.add_paragraph(content.get("adjustment", "教师可在此补充个体化调整建议"))

    try:
        doc.core_properties.comments = (
            f"智伴幼师导出 · v{APP_VERSION} · 排版引擎 python-docx（观察记录）"
        )
    except Exception:
        pass
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


def _fill_daily_template_docx_bytes(cleaned_bytes: bytes, fill_data: dict[str, str]) -> bytes:
    doc = Document(io.BytesIO(cleaned_bytes))
    filled_tc_elems: set = set()
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            cells = row.cells
            for col_idx, cell in enumerate(cells):
                field = match_daily_field(_get_cell_text(cell))
                if not field:
                    continue
                content = fill_data.get(field)
                if not content:
                    continue
                target = None
                if col_idx + 1 < len(cells):
                    cand = cells[col_idx + 1]
                    if not match_daily_field(_get_cell_text(cand)) and not _is_colored_cell(cand):
                        target = cand
                if target is None and row_idx + 1 < len(table.rows) and col_idx < len(table.rows[row_idx + 1].cells):
                    cand = table.rows[row_idx + 1].cells[col_idx]
                    if not match_daily_field(_get_cell_text(cand)) and not _is_colored_cell(cand):
                        target = cand
                if target is None:
                    continue
                tc_elem = target._tc
                if tc_elem in filled_tc_elems:
                    continue
                filled_tc_elems.add(tc_elem)
                _write_cell_preserve_style(target, content)
    try:
        doc.core_properties.comments = (
            f"智伴幼师导出 · v{APP_VERSION} · 排版引擎 python-docx（Aspose 回退）"
        )
    except Exception:
        pass
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


def _fill_daily_template_aspose_bytes(cleaned_bytes: bytes, fill_data: dict[str, str]) -> bytes:
    _aw_require()
    doc = _aw_doc_from_bytes(cleaned_bytes)
    filled_cells: set[int] = set()
    tables = doc.get_child_nodes(aw.NodeType.TABLE, True)
    for ti in range(tables.count):
        table = _aw_node_as_table(tables[ti])
        for row_idx in range(table.rows.count):
            row = table.rows[row_idx]
            for col_idx in range(row.cells.count):
                cell = row.cells[col_idx]
                field = match_daily_field(_aw_cell_text(cell))
                if not field:
                    continue
                content = fill_data.get(field)
                if not content:
                    continue
                target = None
                if col_idx + 1 < row.cells.count:
                    cand = row.cells[col_idx + 1]
                    if not match_daily_field(_aw_cell_text(cand)) and not _aw_cell_has_color(cand):
                        target = cand
                if target is None and row_idx + 1 < table.rows.count and col_idx < table.rows[row_idx + 1].cells.count:
                    cand = table.rows[row_idx + 1].cells[col_idx]
                    if not match_daily_field(_aw_cell_text(cand)) and not _aw_cell_has_color(cand):
                        target = cand
                if target is None:
                    continue
                key = _aw_cell_dedup_key(target)
                if key in filled_cells:
                    continue
                filled_cells.add(key)
                _aw_write_cell_preserve_style(doc, target, content)
    _aw_stamp_export_provenance(doc)
    return _aw_doc_to_bytes(doc)


# ── 日教案 Word 铁律填充 ────────────────────────────────────────────
def fill_daily_word_template(
    template_bytes: bytes,
    daily_content: dict,
    week_theme: str,
    day: str,
    phil: str,
) -> tuple[bytes, str]:
    """日教案填充：Aspose 优先，失败回退 python-docx。返回 (docx 字节, 引擎标识)。"""
    cleaned_bytes = clean_template_keep_style(template_bytes)
    fill_data = _build_daily_fill_data(daily_content, week_theme, day, phil)
    try:
        if ENABLE_ASPOSE_WORDS:
            return (
                _fill_daily_template_aspose_bytes(cleaned_bytes, fill_data),
                "aspose-words",
            )
    except Exception as e:
        logger.warning("Aspose 日教案填充失败，回退 python-docx：%s", e, exc_info=True)
    return (
        _fill_daily_template_docx_bytes(cleaned_bytes, fill_data),
        "python-docx",
    )


# ── /generate-weekly ───────────────────────────────────────────────
@app.post("/generate-weekly", tags=["周日联动"])
async def generate_weekly(
    theme:      str = Form(..., description="周主题"),
    phil:       str = Form(..., description="教育理念"),
    activities: str = Form("[]", description="活动类型列表（JSON）"),
    class_level: str = Form("中班", description="班级（小班/中班/大班）"),
):
    """
    基于主题、理念和班级，自动生成高质量周计划。

    使用 Prompt 工程系统确保输出风格一致、质量稳定（可复现）。

    返回 JSON，包含每天的详细活动信息。
    前端可直接展示，并允许用户选择某一天生成日教案。
    """
    try:
        acts_list: list[str] = json.loads(activities) if activities else []
        if not isinstance(acts_list, list):
            acts_list = [str(acts_list)]
    except (json.JSONDecodeError, TypeError):
        # 支持直接传字符串，如 "户外活动" 或 "户外活动,区域活动"
        acts_list = [a.strip() for a in activities.split(",") if a.strip()] if activities else []
    plan = generate_weekly_content(theme, phil, acts_list, class_level)
    return {"status": "ok", "weekly_plan": plan}


@app.post("/generate-term-plan", tags=["园部计划"])
async def generate_term_plan(
    term_theme: str = Form(..., description="学期主题"),
    start_month: int = Form(2, description="起始月份（1-12）"),
    month_count: int = Form(5, description="学期月数（1-6）"),
):
    """
    园部计划骨架生成：输出学期→月→周三级结构。
    用于 To B 端先统一园所节奏，再分发到班级周计划。
    """
    if not (1 <= start_month <= 12):
        raise HTTPException(status_code=400, detail="start_month 必须在 1-12 之间")
    skeleton = _build_term_month_week_skeleton(term_theme, start_month, month_count)
    return {"status": "ok", "term_plan": skeleton}


@app.post("/apply-daily-feedback", tags=["周日联动"])
async def apply_daily_feedback(
    weekly_plan: str = Form(..., description="周计划 JSON"),
    day: str = Form(..., description="目标星期，如 周一"),
    completion_score: int = Form(..., description="执行完成度（1-5）"),
    highlights: str = Form("", description="今日亮点"),
    risks: str = Form("", description="风险与问题"),
    adjust_suggestion: str = Form("", description="下次调整建议"),
):
    """
    日计划执行后反馈回流周计划：
    将某一天的完成度与复盘意见写回 weekly_plan，形成可迭代闭环。
    """
    if completion_score < 1 or completion_score > 5:
        raise HTTPException(status_code=400, detail="completion_score 必须在 1-5 之间")
    try:
        plan: dict = json.loads(weekly_plan)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="weekly_plan 不是合法 JSON")

    days: list[dict] = plan.get("days", [])
    target_idx = next((i for i, d in enumerate(days) if d.get("day") == day), None)
    if target_idx is None:
        raise HTTPException(status_code=400, detail=f"未找到目标日期：{day}")

    day_item = copy.deepcopy(days[target_idx])
    day_item["execution_feedback"] = {
        "completion_score": completion_score,
        "highlights": highlights.strip(),
        "risks": risks.strip(),
        "adjust_suggestion": adjust_suggestion.strip(),
    }
    day_item["status"] = "reviewed"
    days[target_idx] = day_item
    plan["days"] = days

    return {
        "status": "ok",
        "updated_weekly_plan": plan,
        "updated_day": day_item,
    }


@app.get("/roadmap", tags=["版本规划"])
async def roadmap():
    """模块分期配置：P0 已落地，P1/P2 预留，供前端展示与后续接口扩展。"""
    return {
        "status": "ok",
        "version": "v1.2.0",
        "tracks": {
            "to_b": [
                {"module": "模块1 园部学期-月-周计划", "phase": "P0", "status": "ready"},
                {"module": "模块4 幼儿成长中台", "phase": "P1", "status": "reserved"},
                {"module": "模块6 家园沟通中台", "phase": "P1", "status": "reserved"},
                {"module": "模块7 教师发展中台", "phase": "P2", "status": "reserved"},
            ],
            "to_c": [
                {"module": "模块2 日计划生成与调整并回流周计划", "phase": "P0", "status": "ready"},
                {"module": "模块3 拍照观察与现场记录", "phase": "P0", "status": "ready"},
                {"module": "模块5 多场景活动引擎", "phase": "P2", "status": "reserved"},
                {"module": "模块7 教师个人档案与成长", "phase": "P2", "status": "reserved"},
            ],
        },
    }


# ── /generate-daily ────────────────────────────────────────────────
@app.post("/generate-daily", tags=["周日联动"])
async def generate_daily(
    weekly_plan: str  = Form(..., description="周计划 JSON（由 /generate-weekly 返回或前端暂存）"),
    day:         str  = Form(..., description="目标星期，如 周一"),
    phil:        str  = Form(..., description="教育理念"),
    template:    Optional[UploadFile] = File(None, description="日教案 Word 模板 (.docx，可选)"),
):
    """
    周→日联动接口：将周计划中某一天的任务拆解为四维日教案并导出 Word。

    四维结构：**导入**（情境激活）→ **过程**（分步操作）→ **延伸**（区域/家园）→ **反思**（观察要点）

    若上传模板：识别模板关键字单元格并回填。
    若未上传模板：自动生成结构化（非表格）日计划文档，并追加「生成」「调整」两部分。
    """
    template_bytes: Optional[bytes] = None
    if template is not None:
        if not (template.filename or "").lower().endswith(".docx"):
            raise HTTPException(status_code=400, detail="仅支持 .docx 格式")
        template_bytes = await template.read()
        if not template_bytes:
            raise HTTPException(status_code=400, detail="上传的文件为空")

    # 解析周计划 JSON
    try:
        plan: dict = json.loads(weekly_plan)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="weekly_plan 不是合法 JSON")

    # 找到目标日的任务
    days: list[dict] = plan.get("days", [])
    target = next((d for d in days if d.get("day") == day), None)
    if not target:
        raise HTTPException(
            status_code=400,
            detail=f"在周计划中未找到「{day}」，可用值：{[d.get('day') for d in days]}",
        )

    week_theme = plan.get("week_theme", "本周主题")
    task       = target.get("task", day)

    # AI 生成四维日教案
    daily_content = generate_daily_content(
        week_theme=week_theme,
        day=day,
        task=task,
        phil=phil,
    )

    if template_bytes:
        try:
            filled_bytes, export_engine = fill_daily_word_template(
                template_bytes=template_bytes,
                daily_content=daily_content,
                week_theme=week_theme,
                day=day,
                phil=phil,
            )
        except PackageNotFoundError:
            raise HTTPException(status_code=400, detail="模板解析失败，请确认上传的是有效 .docx 文件")
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"日教案模板填充失败：{e}")
        original_name = os.path.basename(template.filename or "daily-template.docx")
        original_name = re.sub(r'[\\/*?:"<>|]+', "_", original_name)
        if not original_name.lower().endswith(".docx"):
            original_name += ".docx"
    else:
        filled_bytes = _build_daily_structured_docx_bytes(
            daily_content=daily_content,
            week_theme=week_theme,
            day=day,
            phil=phil,
            source_day=target,
        )
        export_engine = "python-docx-structured"
        original_name = re.sub(r'[\\/*?:"<>|]+', "_", f"{week_theme}_{day}_日计划.docx")

    h = {"Content-Disposition": _build_content_disposition(original_name)}
    h.update(_export_http_headers(export_engine))
    return StreamingResponse(
        io.BytesIO(filled_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=h,
    )


# ── /preview-daily（调试用，不需要上传文件）──────────────────────────
@app.post("/preview-daily", tags=["调试"])
async def preview_daily(
    weekly_plan: str = Form(...),
    day:         str = Form(...),
    phil:        str = Form(...),
):
    """返回日教案 AI 内容 JSON，用于前端预览调试"""
    plan   = json.loads(weekly_plan)
    days   = plan.get("days", [])
    target = next((d for d in days if d.get("day") == day), None)
    if not target:
        raise HTTPException(status_code=400, detail=f"未找到「{day}」")
    content = generate_daily_content(
        week_theme=plan.get("week_theme", ""),
        day=day,
        task=target.get("task", ""),
        phil=phil,
    )
    return {"status": "ok", "day": day, "task": target.get("task"), "content": content}


# ── /preview-observation ───────────────────────────────────────────
@app.post("/preview-observation", tags=["拍照观察"])
async def preview_observation(
    theme: str = Form(...),
    child_name: str = Form(""),
    scene: str = Form("活动现场"),
    note: str = Form(""),
    phil: str = Form("五大领域"),
    photo_names: str = Form("[]"),
):
    try:
        parsed = json.loads(photo_names) if photo_names else []
        names = [str(x) for x in parsed if str(x).strip()]
    except Exception:
        names = []
    content = generate_observation_content(
        theme=theme,
        child_name=child_name,
        scene=scene,
        note=note,
        phil=phil,
        photo_names=names,
    )
    return {"status": "ok", "content": content}


# ── /generate-observation ──────────────────────────────────────────
@app.post("/generate-observation", tags=["拍照观察"])
async def generate_observation(
    theme: str = Form(...),
    child_name: str = Form(""),
    scene: str = Form("活动现场"),
    note: str = Form(""),
    phil: str = Form("五大领域"),
    photos: list[UploadFile] = File(default=[]),
):
    valid_photos = [p for p in photos if (p.filename or "").strip()]
    photo_names = [str(p.filename).strip() for p in valid_photos]
    content = generate_observation_content(
        theme=theme,
        child_name=child_name,
        scene=scene,
        note=note,
        phil=phil,
        photo_names=photo_names,
    )
    filled_bytes = _build_observation_docx_bytes(
        content=content,
        theme=theme,
        child_name=child_name,
        scene=scene,
        phil=phil,
        note=note,
        photo_names=photo_names,
    )
    child = child_name.strip() or "幼儿"
    original_name = re.sub(r'[\\/*?:"<>|]+', "_", f"{theme}_{child}_观察记录.docx")
    h = {"Content-Disposition": _build_content_disposition(original_name)}
    h.update(_export_http_headers("python-docx-observation"))
    return StreamingResponse(
        io.BytesIO(filled_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=h,
    )


@app.post("/generate-mini", tags=["核心接口"])
async def generate_mini(
    theme: str = Form(..., description="教学主题"),
    template: UploadFile = File(..., description="Word 模板文件 (.doc/.docx)"),
    phil: str = Form("五大领域", description="教育理念"),
    activities: str = Form("[]", description="活动重点列表（JSON 数组字符串）"),
    child_initiative: bool = Form(False, description="是否有幼儿自主发起活动"),
    child_desc: str = Form("", description="幼儿自主活动描述"),
    class_level: str = Form("", description="班级类型：小班 / 中班 / 大班"),
):
    """
    小程序友好版生成接口。

    负责接收老师上传的模板，生成后先落盘，再返回一个可下载的临时链接。
    """
    if not (template.filename or "").lower().endswith(".docx"):
        raise HTTPException(
            status_code=400,
            detail="仅支持 .docx 格式的 Word 文件",
        )

    try:
        acts_list: list[str] = json.loads(activities) if activities else []
    except (json.JSONDecodeError, TypeError):
        acts_list = []
    if not isinstance(acts_list, list):
        acts_list = []

    template_bytes = await template.read()
    if len(template_bytes) == 0:
        raise HTTPException(status_code=400, detail="上传的文件为空")

    _append_jsonl(_WEEKLY_DRAFT_LOG_FILE, {
        "ts": _utc_iso(),
        "event": "template_submit",
        "flow": "mini_template_generate",
        "template_filename": os.path.basename(template.filename or "template.docx"),
        "theme": theme,
        "phil": phil,
        "activities": acts_list,
        "child_initiative": bool(child_initiative),
    })

    template_outline = _extract_template_outline(template_bytes)
    ai_content = generate_content(
        theme=theme,
        phil=phil,
        activities=acts_list,
        child_initiative=child_initiative,
        child_desc=child_desc,
        template_outline=template_outline,
        class_level=class_level.strip(),
    )

    try:
        filled_bytes, export_engine = fill_word_template(
            template_bytes=template_bytes,
            theme=theme,
            phil=phil,
            ai_content=ai_content,
            activities=acts_list,
            child_initiative=child_initiative,
            child_desc=child_desc,
            class_level=class_level.strip(),
        )
    except PackageNotFoundError:
        raise HTTPException(status_code=400, detail="模板解析失败，请确认上传的是有效 .docx 文件")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"模板填充失败：{e}")

    original_name = os.path.basename(template.filename or "template.docx")
    original_name = re.sub(r'[\\/*?:"<>|]+', "_", original_name)
    if not original_name.lower().endswith(".docx"):
        original_name += ".docx"

    return _build_mini_doc_payload(
        filled_bytes=filled_bytes,
        original_name=original_name,
        export_engine=export_engine,
    )


@app.post("/mini-template/upload", tags=["核心接口"])
async def upload_mini_template(
    template: UploadFile = File(..., description="Word 模板文件 (.docx)"),
):
    """
    小程序模板上传接口。

    - 先把老师模板保存到云端临时目录
    - 返回 template_id，后续生成只需传 template_id
    """
    original_name = os.path.basename(template.filename or "template.docx")
    if not original_name.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="仅支持 .docx 格式的 Word 文件")

    template_bytes = await template.read()
    if not template_bytes:
        raise HTTPException(status_code=400, detail="上传的文件为空")

    template_id = uuid4().hex
    template_path = _TEMP_TEMPLATE_DIR / f"{template_id}.docx"
    template_path.write_bytes(template_bytes)

    _TEMP_TEMPLATES[template_id] = {
        "path": str(template_path),
        "filename": original_name,
        "size": str(len(template_bytes)),
        "uploaded_at_utc": _utc_iso(),
    }
    return {
        "status": "ok",
        "template_id": template_id,
        "filename": original_name,
        "size": len(template_bytes),
    }


@app.post("/generate-mini-by-template", tags=["核心接口"])
async def generate_mini_by_template(
    theme: str = Form(..., description="教学主题"),
    template_id: str = Form(..., description="模板ID（由 /mini-template/upload 返回）"),
    phil: str = Form("五大领域", description="教育理念"),
    activities: str = Form("[]", description="活动重点列表（JSON 数组字符串）"),
    child_initiative: bool = Form(False, description="是否有幼儿自主发起活动"),
    child_desc: str = Form("", description="幼儿自主活动描述"),
    class_level: str = Form("", description="班级类型：小班 / 中班 / 大班"),
):
    """
    小程序按 template_id 生成接口。

    先从云端临时模板池读取模板，再执行和 /generate-mini 一致的生成流程。
    """
    temp_info = _TEMP_TEMPLATES.get((template_id or "").strip())
    if not temp_info:
        raise HTTPException(status_code=404, detail="模板不存在或已过期，请重新上传")

    template_path = Path(temp_info.get("path", ""))
    if not template_path.exists():
        raise HTTPException(status_code=404, detail="模板不存在或已过期，请重新上传")

    try:
        acts_list: list[str] = json.loads(activities) if activities else []
    except (json.JSONDecodeError, TypeError):
        acts_list = []
    if not isinstance(acts_list, list):
        acts_list = []

    template_bytes = template_path.read_bytes()
    if len(template_bytes) == 0:
        raise HTTPException(status_code=400, detail="模板文件为空，请重新上传")

    _append_jsonl(_WEEKLY_DRAFT_LOG_FILE, {
        "ts": _utc_iso(),
        "event": "template_submit",
        "flow": "mini_template_generate_by_id",
        "template_id": template_id,
        "template_filename": temp_info.get("filename") or "template.docx",
        "theme": theme,
        "phil": phil,
        "activities": acts_list,
        "child_initiative": bool(child_initiative),
    })

    template_outline = _extract_template_outline(template_bytes)
    ai_content = generate_content(
        theme=theme,
        phil=phil,
        activities=acts_list,
        child_initiative=child_initiative,
        child_desc=child_desc,
        template_outline=template_outline,
        class_level=class_level.strip(),
    )

    try:
        filled_bytes, export_engine = fill_word_template(
            template_bytes=template_bytes,
            theme=theme,
            phil=phil,
            ai_content=ai_content,
            activities=acts_list,
            child_initiative=child_initiative,
            child_desc=child_desc,
            class_level=class_level.strip(),
        )
    except PackageNotFoundError:
        raise HTTPException(status_code=400, detail="模板解析失败，请确认上传的是有效 .docx 文件")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"模板填充失败：{e}")

    original_name = os.path.basename(temp_info.get("filename") or "template.docx")
    original_name = re.sub(r'[\\/*?:"<>|]+', "_", original_name)
    if not original_name.lower().endswith(".docx"):
        original_name += ".docx"

    return _build_mini_doc_payload(
        filled_bytes=filled_bytes,
        original_name=original_name,
        export_engine=export_engine,
    )


@app.get("/mini-export/{token}", tags=["核心接口"])
async def download_mini_export(token: str):
    info = _TEMP_EXPORTS.get(token)
    if not info:
        raise HTTPException(status_code=404, detail="文件不存在或已过期")

    path = Path(info["path"])
    if not path.exists():
        raise HTTPException(status_code=404, detail="文件不存在或已过期")

    headers = {
        "Content-Disposition": _build_content_disposition(info["filename"]),
    }
    headers.update(_export_http_headers(info.get("engine", "python-docx-mini")))
    return FileResponse(
        path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=info["filename"],
        headers=headers,
    )


@app.post("/generate-observation-mini", tags=["拍照观察"])
async def generate_observation_mini(
    theme: str = Form(...),
    child_name: str = Form(""),
    scene: str = Form("活动现场"),
    note: str = Form(""),
    phil: str = Form("五大领域"),
    photo_names: str = Form("[]"),
):
    """
    小程序友好版观察记录生成接口。

    仅依赖文本和照片名列表，避免小程序端多文件上传的额外复杂度。
    """
    try:
        parsed = json.loads(photo_names) if photo_names else []
        names = [str(x) for x in parsed if str(x).strip()]
    except Exception:
        names = []
    content = generate_observation_content(
        theme=theme,
        child_name=child_name,
        scene=scene,
        note=note,
        phil=phil,
        photo_names=names,
    )
    filled_bytes = _build_observation_docx_bytes(
        content=content,
        theme=theme,
        child_name=child_name,
        scene=scene,
        phil=phil,
        note=note,
        photo_names=names,
    )
    child = child_name.strip() or "幼儿"
    original_name = re.sub(r'[\\/*?:"<>|]+', "_", f"{theme}_{child}_观察记录.docx")
    h = {"Content-Disposition": _build_content_disposition(original_name)}
    h.update(_export_http_headers("python-docx-observation-mini"))
    return StreamingResponse(
        io.BytesIO(filled_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=h,
    )


@app.post("/transcribe-voice-mini", tags=["语音输入"])
async def transcribe_voice_mini(
    audio: UploadFile = File(..., description="录音文件"),
    prompt: str = Form("", description="转写提示词"),
):
    """
    小程序语音输入转文字接口。

    当前优先使用 OpenAI 语音转文字能力；未配置 OPENAI_API_KEY 时返回可读错误。
    """
    if voice_client is None:
        raise HTTPException(
            status_code=501,
            detail="语音转文字未配置：请先设置 OPENAI_API_KEY",
        )

    original_name = os.path.basename(audio.filename or "voice.mp3")
    suffix = Path(original_name).suffix or ".mp3"
    audio_bytes = await audio.read()
    if not audio_bytes:
        raise HTTPException(status_code=400, detail="录音文件为空")

    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(audio_bytes)
        tmp_path = Path(tmp.name)

    try:
        with tmp_path.open("rb") as audio_file:
            transcription = voice_client.audio.transcriptions.create(
                model=VOICE_TRANSCRIBE_MODEL,
                file=audio_file,
                prompt=prompt.strip() or None,
            )
        text = getattr(transcription, "text", "") or ""
        text = text.strip()
        if not text:
            raise HTTPException(status_code=502, detail="语音识别成功，但未返回文本")
        return {
            "status": "ok",
            "text": text,
            "model": VOICE_TRANSCRIBE_MODEL,
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"语音转文字失败：{e}")
    finally:
        try:
            tmp_path.unlink(missing_ok=True)
        except Exception:
            pass


# ──────────────────────────────────────────────
# 第三方对接测试页（静态）
# ──────────────────────────────────────────────
@app.get("/mock-token", tags=["工具"], include_in_schema=False)
async def mock_token_page():
    """第三方商城对接测试台（临时测试页）。"""
    p = _BASE_DIR / "mock-mall.html"
    if p.exists():
        return FileResponse(str(p), media_type="text/html")
    raise HTTPException(status_code=404, detail="测试页未找到")


# ──────────────────────────────────────────────
# C. 批量速写接口（asyncio 并发，每位老师用自己的 agent）
# ──────────────────────────────────────────────

async def _generate_weekly_for_user(
    openid: str,
    theme: str,
    phil: str,
    activities: list[str],
    class_level: str,
) -> dict:
    """单用户异步生成周计划，注入该用户的 agent_profile。"""
    try:
        accounts = _load_user_accounts()
        entry = accounts.get(openid, {})
        agent = entry.get("agent_profile", {})

        # 把 agent 性格注入到 prompt 里
        prompt_template = get_prompt_template()
        system_prompt = prompt_template.build_system_prompt()
        if agent:
            agent_hint = (
                f"\n\n【本次老师的智能体设定】\n"
                f"名字：{agent.get('name','小助手')}\n"
                f"性格：{agent.get('personality','热心、耐心')}\n"
                f"音调：{agent.get('tone','亲切温暖')}\n"
                f"教学风格：{agent.get('style','鼓励式教学')}\n"
                f"请按照以上风格生成内容，让老师感受到专属感。"
            )
            system_prompt = system_prompt + agent_hint

        user_prompt = prompt_template.build_user_prompt(
            theme=theme,
            class_level=class_level or "中班",
            philosophy=phil,
            activities=activities or ["区域活动", "户外活动"],
        )

        resp = client.chat.completions.create(
            model=AI_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.2,
            max_tokens=4096,
        )
        raw = resp.choices[0].message.content.strip()
        plan = _parse_json_payload(raw)

        # 写入 Firestore 历史
        if FIRESTORE_ENABLED:
            try:
                _fs().collection("users").document(openid)\
                    .collection("history").add({
                        "type": "weekly_plan",
                        "theme": theme,
                        "plan": plan,
                        "created_at": _utc_iso(),
                    })
            except Exception as e:
                logger.warning("Firestore 写历史失败：%s", e)

        return {"openid": openid, "ok": True, "plan": plan}
    except Exception as e:
        logger.error("批量生成失败 openid=%s err=%s", openid, e)
        return {"openid": openid, "ok": False, "error": str(e)}


@app.post("/batch/generate-weekly", tags=["批量速写"])
async def batch_generate_weekly(payload: dict = Body(...)):
    """
    批量速写接口：同时为多位老师生成个性化周计划。

    每位老师自动使用自己的 agent_profile（性格/风格），asyncio 并发执行。

    请求体：
    {
      "openids": ["openid_1", "openid_2"],   # 目标老师列表
      "theme": "春天来了",
      "phil": "以幼儿为中心",
      "activities": ["户外活动", "区域活动"],
      "class_level": "中班",
      "admin_token": "..."                   # 管理员 token（可选鉴权）
    }
    """
    openids: list[str] = payload.get("openids", [])
    theme = str(payload.get("theme", "")).strip()
    phil = str(payload.get("phil", "以幼儿为中心")).strip()
    activities = payload.get("activities", ["户外活动", "区域活动"])
    class_level = str(payload.get("class_level", "中班")).strip()

    if not openids:
        raise HTTPException(status_code=400, detail="openids 不能为空")
    if not theme:
        raise HTTPException(status_code=400, detail="theme 不能为空")
    if len(openids) > 50:
        raise HTTPException(status_code=400, detail="单次最多 50 位老师")

    # asyncio 并发全部用户
    tasks = [
        _generate_weekly_for_user(openid, theme, phil, activities, class_level)
        for openid in openids
    ]
    results = await asyncio.gather(*tasks)

    ok_count = sum(1 for r in results if r.get("ok"))
    fail_count = len(results) - ok_count

    return {
        "status": "ok",
        "total": len(openids),
        "ok_count": ok_count,
        "fail_count": fail_count,
        "results": list(results),
    }


@app.get("/batch/users", tags=["批量速写"])
async def batch_list_users():
    """列出所有注册用户（openid + agent_profile），用于批量速写前选择目标。"""
    accounts = _load_user_accounts()
    users = []
    for openid, entry in accounts.items():
        users.append({
            "openid": openid,
            "user_id": entry.get("user_id") or "",
            "agent": entry.get("agent_profile", {}),
            "created_at": entry.get("created_at", ""),
        })
    return {"status": "ok", "count": len(users), "users": users}


# ──────────────────────────────────────────────
# 应用启动事件
# ──────────────────────────────────────────────
@app.on_event("startup")
async def _startup() -> None:
    asyncio.create_task(_webhook_retry_loop())
    logger.info("Webhook 重试任务已启动（每 2 分钟检查一次）")


# ──────────────────────────────────────────────
# 本地启动入口
# Cloud Run 通过 PORT 环境变量动态注入端口（默认 8080）
# 本地开发默认使用 8000
# ──────────────────────────────────────────────
if __name__ == "__main__":
    import uvicorn

    port = int(os.environ.get("PORT", 8000))
    # reload=True 仅在本地开发时使用；Cloud Run 不走此分支
    is_local = port == 8000
    uvicorn.run("main:app", host="0.0.0.0", port=port, reload=is_local)
